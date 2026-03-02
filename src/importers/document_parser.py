"""Document parsing utilities for text extraction.

Supports:
- PDF files (via pdfplumber)
- DOCX files (via python-docx)
- Plain text files
- Images with OCR (via pytesseract) - optional
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field
from datetime import datetime

# PDF parsing
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# OCR support (optional)
try:
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    file_path: str
    file_name: str
    file_type: str
    text: str
    page_count: int = 1
    word_count: int = 0
    sections: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    parse_errors: List[str] = field(default_factory=list)
    parsed_at: str = field(default_factory=lambda: datetime.now().isoformat())

    def to_dict(self) -> Dict[str, Any]:
        return {
            "file_path": self.file_path,
            "file_name": self.file_name,
            "file_type": self.file_type,
            "text": self.text,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "sections": self.sections,
            "metadata": self.metadata,
            "parse_errors": self.parse_errors,
            "parsed_at": self.parsed_at,
        }


class DocumentParser:
    """Parse documents and extract text content."""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'text',
        '.md': 'text',
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.tiff': 'image',
        '.bmp': 'image',
    }

    def __init__(self):
        """Initialize parser with available backends."""
        self.pdf_available = PDF_SUPPORT
        self.docx_available = DOCX_SUPPORT
        self.ocr_available = OCR_SUPPORT

    def get_capabilities(self) -> Dict[str, bool]:
        """Return available parsing capabilities."""
        return {
            "pdf": self.pdf_available,
            "docx": self.docx_available,
            "ocr": self.ocr_available,
            "text": True,
        }

    def can_parse(self, file_path: str) -> bool:
        """Check if file type is supported."""
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False

        file_type = self.SUPPORTED_EXTENSIONS[ext]
        if file_type == 'pdf' and not self.pdf_available:
            return False
        if file_type == 'docx' and not self.docx_available:
            return False
        if file_type == 'image' and not self.ocr_available:
            return False

        return True

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document and extract text.

        Args:
            file_path: Path to the document file.

        Returns:
            ParsedDocument with extracted text and metadata.
        """
        path = Path(file_path)

        if not path.exists():
            return ParsedDocument(
                file_path=file_path,
                file_name=path.name,
                file_type="unknown",
                text="",
                parse_errors=[f"File not found: {file_path}"],
            )

        ext = path.suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(ext, 'unknown')

        if file_type == 'pdf':
            return self._parse_pdf(path)
        elif file_type == 'docx':
            return self._parse_docx(path)
        elif file_type == 'text':
            return self._parse_text(path)
        elif file_type == 'image':
            return self._parse_image(path)
        else:
            return ParsedDocument(
                file_path=file_path,
                file_name=path.name,
                file_type=file_type,
                text="",
                parse_errors=[f"Unsupported file type: {ext}"],
            )

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Parse PDF document."""
        if not self.pdf_available:
            return ParsedDocument(
                file_path=str(path),
                file_name=path.name,
                file_type="pdf",
                text="",
                parse_errors=["PDF support not available. Install pdfplumber."],
            )

        errors = []
        pages_text = []
        sections = []
        metadata = {}

        try:
            with pdfplumber.open(path) as pdf:
                metadata = {
                    "page_count": len(pdf.pages),
                    "pdf_info": pdf.metadata or {},
                }

                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text() or ""
                        pages_text.append(text)
                        sections.append({
                            "type": "page",
                            "index": i + 1,
                            "text": text,
                            "char_count": len(text),
                        })
                    except Exception as e:
                        errors.append(f"Error on page {i + 1}: {str(e)}")
                        pages_text.append("")

        except Exception as e:
            errors.append(f"PDF parse error: {str(e)}")

        full_text = "\n\n".join(pages_text)

        return ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            file_type="pdf",
            text=full_text,
            page_count=len(pages_text),
            word_count=len(full_text.split()),
            sections=sections,
            metadata=metadata,
            parse_errors=errors,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Parse DOCX document."""
        if not self.docx_available:
            return ParsedDocument(
                file_path=str(path),
                file_name=path.name,
                file_type="docx",
                text="",
                parse_errors=["DOCX support not available. Install python-docx."],
            )

        errors = []
        paragraphs = []
        sections = []
        metadata = {}

        try:
            doc = DocxDocument(path)

            # Extract core properties if available
            try:
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                }
            except Exception:
                pass

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    sections.append({
                        "type": "paragraph",
                        "index": i,
                        "text": text,
                        "style": para.style.name if para.style else "Normal",
                    })

        except Exception as e:
            errors.append(f"DOCX parse error: {str(e)}")

        full_text = "\n\n".join(paragraphs)

        return ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            file_type="docx",
            text=full_text,
            page_count=1,  # DOCX doesn't have pages in the same way
            word_count=len(full_text.split()),
            sections=sections,
            metadata=metadata,
            parse_errors=errors,
        )

    def _parse_text(self, path: Path) -> ParsedDocument:
        """Parse plain text file."""
        errors = []
        text = ""

        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                text = path.read_text(encoding='latin-1')
        except Exception as e:
            errors.append(f"Text read error: {str(e)}")

        # Split into sections by double newlines
        sections = []
        for i, chunk in enumerate(text.split('\n\n')):
            chunk = chunk.strip()
            if chunk:
                sections.append({
                    "type": "section",
                    "index": i,
                    "text": chunk,
                })

        return ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            file_type="text",
            text=text,
            page_count=1,
            word_count=len(text.split()),
            sections=sections,
            metadata={},
            parse_errors=errors,
        )

    def _parse_image(self, path: Path) -> ParsedDocument:
        """Parse image with OCR."""
        if not self.ocr_available:
            return ParsedDocument(
                file_path=str(path),
                file_name=path.name,
                file_type="image",
                text="",
                parse_errors=["OCR not available. Install pytesseract and PIL."],
            )

        errors = []
        text = ""

        try:
            image = Image.open(path)
            text = pytesseract.image_to_string(image)
        except Exception as e:
            errors.append(f"OCR error: {str(e)}")

        return ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            file_type="image",
            text=text,
            page_count=1,
            word_count=len(text.split()),
            sections=[{"type": "ocr", "index": 0, "text": text}] if text else [],
            metadata={"ocr_engine": "tesseract"},
            parse_errors=errors,
        )


# Module-level parser instance
_parser: Optional[DocumentParser] = None


def get_parser() -> DocumentParser:
    """Get or create the document parser singleton."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


def parse_document(file_path: str) -> ParsedDocument:
    """Convenience function to parse a document.

    Args:
        file_path: Path to the document.

    Returns:
        ParsedDocument with extracted text.
    """
    return get_parser().parse(file_path)
