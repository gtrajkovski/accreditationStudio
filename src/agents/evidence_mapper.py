"""Evidence Mapper Agent.

Builds the critical "Standard → Evidence" mappings used in self-studies,
responses, and compliance documentation. This is the core work accreditation
teams spend months doing manually.
"""

import csv
import io
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Any, List, Optional

from src.agents.base_agent import BaseAgent, AgentType
from src.agents.registry import register_agent
from src.config import Config
from src.core.models import (
    AgentResult,
    EvidenceMap,
    EvidenceMapping,
    EvidenceGap,
    CrosswalkEntry,
    generate_id,
    now_iso,
)
# Lazy imports to avoid dependency issues during testing
# from src.core.standards_store import get_standards_store
# from src.search import get_search_service


@register_agent(AgentType.EVIDENCE_MAPPER)
class EvidenceMapperAgent(BaseAgent):
    """Evidence Mapper Agent.

    Responsibilities:
    - For each standard requirement, fetch candidate evidence from doc index
    - Rank evidence quality and relevance
    - Identify gaps where evidence is missing
    - Propose exhibit labels
    - Generate crosswalk tables for self-studies

    Outputs:
    - evidence_map.json
    - crosswalk_table.csv/docx
    """

    @property
    def agent_type(self) -> AgentType:
        return AgentType.EVIDENCE_MAPPER

    @property
    def system_prompt(self) -> str:
        return """You are the Evidence Mapper Agent for AccreditAI.

Your critical mission is to build "Standard → Evidence" mappings that connect
accreditation requirements to documentary evidence in the institution's documents.

For each standard requirement you must:
1. Search the document index for relevant text passages
2. Rank evidence by quality and directness
3. Identify gaps where evidence is weak or missing
4. Propose exhibit labels (e.g., "Exhibit 1.A - Refund Policy")
5. Note page numbers and section references

Your output format should include:
- Standard citation (e.g., "ACCSC Section VII.A.4")
- Requirement text summary
- Evidence found (with document, page, confidence)
- Evidence quality rating (strong/adequate/weak/missing)
- Suggested exhibit label

Always be precise with citations. Never claim evidence exists if you cannot cite it."""

    @property
    def tools(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "search_evidence",
                "description": "Search document index for evidence matching a requirement",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "requirement_text": {
                            "type": "string",
                            "description": "The requirement text to find evidence for"
                        },
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID to search within"
                        },
                        "doc_types": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Document types to search (optional filter)"
                        },
                        "n_results": {
                            "type": "integer",
                            "description": "Number of results to return",
                            "default": 10
                        }
                    },
                    "required": ["requirement_text", "institution_id"]
                }
            },
            {
                "name": "map_standard_to_evidence",
                "description": "Create an evidence mapping for a specific standard",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "standard_id": {
                            "type": "string",
                            "description": "Standard section ID"
                        },
                        "standard_text": {
                            "type": "string",
                            "description": "Full text of the standard requirement"
                        },
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        }
                    },
                    "required": ["standard_id", "standard_text", "institution_id"]
                }
            },
            {
                "name": "generate_crosswalk_table",
                "description": "Generate a crosswalk table mapping all standards to evidence",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "standards_id": {
                            "type": "string",
                            "description": "Standards library ID"
                        },
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "output_format": {
                            "type": "string",
                            "enum": ["json", "csv", "docx"],
                            "description": "Output format for crosswalk table",
                            "default": "json"
                        },
                        "min_confidence": {
                            "type": "number",
                            "description": "Minimum confidence threshold (0-1)",
                            "default": 0.5
                        }
                    },
                    "required": ["standards_id", "institution_id"]
                }
            },
            {
                "name": "identify_evidence_gaps",
                "description": "Identify standards with weak or missing evidence",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "standards_id": {
                            "type": "string",
                            "description": "Standards library ID"
                        },
                        "evidence_map_id": {
                            "type": "string",
                            "description": "ID of existing evidence map to analyze (optional)"
                        },
                        "threshold": {
                            "type": "number",
                            "description": "Confidence threshold below which to flag gaps",
                            "default": 0.7
                        }
                    },
                    "required": ["institution_id", "standards_id"]
                }
            },
            {
                "name": "save_evidence_map",
                "description": "Save evidence map to workspace",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "evidence_map": {
                            "type": "object",
                            "description": "The evidence map data to save"
                        },
                        "filename": {
                            "type": "string",
                            "description": "Filename for the evidence map"
                        }
                    },
                    "required": ["institution_id", "evidence_map"]
                }
            },
            {
                "name": "get_evidence_summary",
                "description": "Get quick overview of evidence status for an institution",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "standards_id": {
                            "type": "string",
                            "description": "Standards library ID (optional)"
                        }
                    },
                    "required": ["institution_id"]
                }
            }
        ]

    def _execute_tool(self, tool_name: str, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Execute an evidence mapper tool."""
        if tool_name == "search_evidence":
            return self._tool_search_evidence(tool_input)
        elif tool_name == "map_standard_to_evidence":
            return self._tool_map_standard_to_evidence(tool_input)
        elif tool_name == "generate_crosswalk_table":
            return self._tool_generate_crosswalk_table(tool_input)
        elif tool_name == "identify_evidence_gaps":
            return self._tool_identify_evidence_gaps(tool_input)
        elif tool_name == "save_evidence_map":
            return self._tool_save_evidence_map(tool_input)
        elif tool_name == "get_evidence_summary":
            return self._tool_get_evidence_summary(tool_input)
        return {"error": f"Unknown tool: {tool_name}"}

    def _tool_search_evidence(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Search for evidence matching a requirement."""
        try:
            from src.search import get_search_service
            institution_id = tool_input.get("institution_id")
            requirement_text = tool_input.get("requirement_text")
            doc_types = tool_input.get("doc_types")
            n_results = tool_input.get("n_results", 10)

            if not institution_id:
                return {"error": "institution_id is required"}
            if not requirement_text:
                return {"error": "requirement_text is required"}

            search_service = get_search_service(institution_id)

            # Search with optional doc_type filter
            if doc_types and len(doc_types) == 1:
                results = search_service.search(
                    query=requirement_text,
                    n_results=n_results,
                    doc_type=doc_types[0]
                )
            else:
                results = search_service.search(
                    query=requirement_text,
                    n_results=n_results
                )

            evidence_items = []
            for result in results:
                evidence_items.append({
                    "chunk_id": result.chunk.id,
                    "document_id": result.chunk.document_id,
                    "page_number": result.chunk.page_number,
                    "section_header": result.chunk.section_header,
                    "text": result.chunk.text_anonymized[:500] if result.chunk.text_anonymized else "",
                    "relevance_score": result.score,
                    "quality": self._assess_evidence_quality(result.score)
                })

            return {
                "success": True,
                "requirement": requirement_text[:100] + "..." if len(requirement_text) > 100 else requirement_text,
                "evidence_found": len(evidence_items),
                "evidence": evidence_items
            }

        except Exception as e:
            return {"error": str(e)}

    def _assess_evidence_quality(self, score: float) -> str:
        """Assess evidence quality based on relevance score."""
        if score >= 0.85:
            return "strong"
        elif score >= 0.70:
            return "adequate"
        elif score >= 0.50:
            return "weak"
        return "insufficient"

    def _tool_map_standard_to_evidence(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Create evidence mapping for a standard."""
        standard_id = tool_input.get("standard_id")
        standard_text = tool_input.get("standard_text")
        institution_id = tool_input.get("institution_id")

        if not all([standard_id, standard_text, institution_id]):
            return {"error": "standard_id, standard_text, and institution_id are required"}

        # Search for evidence
        search_result = self._tool_search_evidence({
            "requirement_text": standard_text,
            "institution_id": institution_id,
            "n_results": 5
        })

        if "error" in search_result:
            return search_result

        # Determine overall evidence status
        evidence_items = search_result.get("evidence", [])
        if not evidence_items:
            status = "missing"
            confidence = 0.0
        else:
            best_quality = max(e["quality"] for e in evidence_items)
            if best_quality == "strong":
                status = "satisfied"
                confidence = 0.9
            elif best_quality == "adequate":
                status = "partial"
                confidence = 0.7
            else:
                status = "weak"
                confidence = 0.4

        mapping = EvidenceMapping(
            standard_id=standard_id,
            standard_number=standard_id,
            standard_text=standard_text,
            status=status,
            confidence=confidence,
            evidence=evidence_items,
            suggested_exhibit=f"Exhibit {standard_id}" if evidence_items else None,
            gap_notes="Evidence missing or insufficient" if status in ["missing", "weak"] else None
        )

        return {
            "success": True,
            "mapping": mapping.to_dict()
        }

    def _tool_generate_crosswalk_table(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Generate crosswalk table for all standards."""
        from src.core.standards_store import get_standards_store
        standards_id = tool_input.get("standards_id")
        institution_id = tool_input.get("institution_id")
        output_format = tool_input.get("output_format", "json")
        min_confidence = tool_input.get("min_confidence", 0.5)

        if not standards_id or not institution_id:
            return {"error": "standards_id and institution_id are required"}

        # Load standards library
        store = get_standards_store()
        library = store.load(standards_id)
        if not library:
            return {"error": f"Standards library not found: {standards_id}"}

        checklist_items = library.checklist_items
        crosswalk_entries = []

        # Map each checklist item
        for item in checklist_items:
            search_result = self._tool_search_evidence({
                "requirement_text": item.description,
                "institution_id": institution_id,
                "n_results": 3
            })

            evidence = search_result.get("evidence", []) if "error" not in search_result else []
            best_evidence = evidence[0] if evidence else None

            entry = CrosswalkEntry(
                standard_ref=item.number,
                section_reference=item.section_reference,
                category=item.category,
                requirement=item.description,
                evidence_found=len(evidence) > 0,
                quality=best_evidence["quality"] if best_evidence else "missing",
                document_id=best_evidence["document_id"] if best_evidence else None,
                page=best_evidence["page_number"] if best_evidence else None,
                snippet=best_evidence["text"][:200] if best_evidence else None,
                confidence=best_evidence["relevance_score"] if best_evidence else 0.0,
                exhibit_label=f"Exhibit {item.number}" if best_evidence else None
            )
            crosswalk_entries.append(entry)

        # Calculate statistics
        total = len(crosswalk_entries)
        with_evidence = sum(1 for e in crosswalk_entries if e.evidence_found)
        stats = {
            "total_standards": total,
            "with_evidence": with_evidence,
            "strong": sum(1 for e in crosswalk_entries if e.quality == "strong"),
            "adequate": sum(1 for e in crosswalk_entries if e.quality == "adequate"),
            "weak": sum(1 for e in crosswalk_entries if e.quality == "weak"),
            "missing": sum(1 for e in crosswalk_entries if e.quality == "missing"),
            "coverage_percent": round(with_evidence / total * 100, 1) if total > 0 else 0
        }

        # Prepare output directory
        output_dir = Config.WORKSPACE_DIR / institution_id / "crossrefs"
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        crosswalk_id = f"xwalk_{timestamp}"

        # Export based on format
        entries_dict = [e.to_dict() for e in crosswalk_entries]

        if output_format == "json":
            output_path = output_dir / f"{crosswalk_id}.json"
            output_data = {
                "id": crosswalk_id,
                "standards_library_id": standards_id,
                "institution_id": institution_id,
                "created_at": now_iso(),
                "statistics": stats,
                "entries": entries_dict
            }
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(output_data, f, indent=2)

        elif output_format == "csv":
            output_path = output_dir / f"{crosswalk_id}.csv"
            with open(output_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=[
                    "standard_ref", "section_reference", "category", "requirement",
                    "evidence_found", "quality", "document_id", "page",
                    "snippet", "confidence", "exhibit_label"
                ])
                writer.writeheader()
                for entry in entries_dict:
                    writer.writerow(entry)

        elif output_format == "docx":
            # DOCX export requires python-docx - create simple fallback
            try:
                from docx import Document
                from docx.shared import Inches

                doc = Document()
                doc.add_heading(f"Evidence Crosswalk - {library.name}", 0)
                doc.add_paragraph(f"Institution: {institution_id}")
                doc.add_paragraph(f"Generated: {now_iso()}")
                doc.add_paragraph(f"Coverage: {stats['coverage_percent']}%")

                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Standard"
                hdr_cells[1].text = "Requirement"
                hdr_cells[2].text = "Quality"
                hdr_cells[3].text = "Document"
                hdr_cells[4].text = "Exhibit"

                for entry in crosswalk_entries:
                    row_cells = table.add_row().cells
                    row_cells[0].text = entry.standard_ref
                    row_cells[1].text = entry.requirement[:100] + "..." if len(entry.requirement) > 100 else entry.requirement
                    row_cells[2].text = entry.quality
                    row_cells[3].text = entry.document_id or "N/A"
                    row_cells[4].text = entry.exhibit_label or "N/A"

                output_path = output_dir / f"{crosswalk_id}.docx"
                doc.save(output_path)
            except ImportError:
                # Fall back to JSON if python-docx not available
                output_format = "json"
                output_path = output_dir / f"{crosswalk_id}.json"
                output_data = {
                    "id": crosswalk_id,
                    "standards_library_id": standards_id,
                    "institution_id": institution_id,
                    "created_at": now_iso(),
                    "statistics": stats,
                    "entries": entries_dict,
                    "note": "DOCX export unavailable, saved as JSON"
                }
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(output_data, f, indent=2)

        return {
            "success": True,
            "crosswalk_id": crosswalk_id,
            "output_path": str(output_path),
            "format": output_format,
            "statistics": stats
        }

    def _tool_identify_evidence_gaps(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Identify gaps in evidence mapping."""
        from src.core.standards_store import get_standards_store
        institution_id = tool_input.get("institution_id")
        standards_id = tool_input.get("standards_id")
        evidence_map_id = tool_input.get("evidence_map_id")
        threshold = tool_input.get("threshold", Config.AGENT_CONFIDENCE_THRESHOLD)

        if not institution_id or not standards_id:
            return {"error": "institution_id and standards_id are required"}

        # Try to load existing evidence map
        evidence_map_data = None
        if evidence_map_id:
            evidence_map_data = self._load_evidence_map(institution_id, evidence_map_id)

        # If no map found, generate fresh crosswalk
        if not evidence_map_data:
            crosswalk_result = self._tool_generate_crosswalk_table({
                "standards_id": standards_id,
                "institution_id": institution_id,
                "output_format": "json"
            })
            if "error" in crosswalk_result:
                return crosswalk_result

            # Load the generated crosswalk
            crosswalk_path = crosswalk_result.get("output_path")
            if crosswalk_path:
                with open(crosswalk_path, "r", encoding="utf-8") as f:
                    evidence_map_data = json.load(f)

        if not evidence_map_data:
            return {"error": "Could not load or generate evidence map"}

        # Load standards for additional context
        store = get_standards_store()
        library = store.load(standards_id)
        checklist_lookup = {}
        if library:
            checklist_lookup = {item.number: item for item in library.checklist_items}

        # Classify gaps by severity
        critical_gaps = []
        high_gaps = []
        advisory_gaps = []

        entries = evidence_map_data.get("entries", [])
        for entry in entries:
            confidence = entry.get("confidence", 0.0)
            quality = entry.get("quality", "missing")

            if confidence < threshold or quality in ["missing", "weak", "insufficient"]:
                standard_ref = entry.get("standard_ref", "")
                checklist_item = checklist_lookup.get(standard_ref)

                gap = EvidenceGap(
                    standard_id=standard_ref,
                    standard_number=standard_ref,
                    standard_text=entry.get("requirement", "")[:200],
                    current_coverage=quality,
                    confidence=confidence,
                    suggestions=self._generate_gap_suggestions(entry, checklist_item),
                    related_doc_types=checklist_item.applies_to if checklist_item else []
                )

                # Classify severity
                category = entry.get("category", "").lower()
                if quality == "missing":
                    if "required" in category or "core" in category:
                        gap.severity = "critical"
                        critical_gaps.append(gap.to_dict())
                    else:
                        gap.severity = "high"
                        high_gaps.append(gap.to_dict())
                else:
                    gap.severity = "advisory"
                    advisory_gaps.append(gap.to_dict())

        return {
            "success": True,
            "total_gaps": len(critical_gaps) + len(high_gaps) + len(advisory_gaps),
            "critical_gaps": critical_gaps,
            "high_gaps": high_gaps,
            "advisory_gaps": advisory_gaps,
            "summary": {
                "critical_count": len(critical_gaps),
                "high_count": len(high_gaps),
                "advisory_count": len(advisory_gaps),
                "threshold_used": threshold
            }
        }

    def _generate_gap_suggestions(
        self,
        entry: Dict[str, Any],
        checklist_item: Optional[Any] = None
    ) -> List[str]:
        """Generate suggestions for filling an evidence gap."""
        suggestions = []
        applies_to = checklist_item.applies_to if checklist_item else []

        if "catalog" in applies_to:
            suggestions.append("Review institutional catalog for relevant content")
        if "policy_manual" in applies_to:
            suggestions.append("Check policy manual for documented procedures")
        if "enrollment_agreement" in applies_to:
            suggestions.append("Verify enrollment agreement includes this requirement")
        if "student_handbook" in applies_to:
            suggestions.append("Check student handbook for relevant policies")
        if "faculty_handbook" in applies_to:
            suggestions.append("Review faculty handbook for applicable content")
        if "self_evaluation_report" in applies_to:
            suggestions.append("Address in self-evaluation report narrative")

        if not suggestions:
            requirement = entry.get("requirement", "this requirement")[:50]
            suggestions.append(f"Upload documents addressing: {requirement}...")

        return suggestions

    def _load_evidence_map(
        self,
        institution_id: str,
        evidence_map_id: str
    ) -> Optional[Dict[str, Any]]:
        """Load an evidence map from workspace."""
        # Check crossrefs directory first
        crossrefs_dir = Config.WORKSPACE_DIR / institution_id / "crossrefs"
        if crossrefs_dir.exists():
            for path in crossrefs_dir.glob(f"{evidence_map_id}*.json"):
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)

        # Check evidence_maps directory
        maps_dir = Config.WORKSPACE_DIR / institution_id / "evidence_maps"
        if maps_dir.exists():
            for path in maps_dir.glob(f"*{evidence_map_id}*.json"):
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)

        return None

    def _tool_save_evidence_map(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Save evidence map to workspace."""
        institution_id = tool_input.get("institution_id")
        evidence_map = tool_input.get("evidence_map")
        filename = tool_input.get("filename", "evidence_map.json")

        if not institution_id or not evidence_map:
            return {"error": "institution_id and evidence_map are required"}

        # Save to workspace
        output_dir = Config.WORKSPACE_DIR / institution_id / "evidence_maps"
        output_dir.mkdir(parents=True, exist_ok=True)

        output_path = output_dir / filename
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(evidence_map, f, indent=2)

        return {
            "success": True,
            "path": str(output_path),
            "message": f"Evidence map saved to {output_path}"
        }

    def _tool_get_evidence_summary(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Get quick evidence status summary."""
        institution_id = tool_input.get("institution_id")
        standards_id = tool_input.get("standards_id")

        if not institution_id:
            return {"error": "institution_id is required"}

        # Check for existing evidence maps
        evidence_maps_dir = Config.WORKSPACE_DIR / institution_id / "evidence_maps"
        crossrefs_dir = Config.WORKSPACE_DIR / institution_id / "crossrefs"

        # Find most recent map from either directory
        map_files = []
        if evidence_maps_dir.exists():
            map_files.extend(evidence_maps_dir.glob("*.json"))
        if crossrefs_dir.exists():
            map_files.extend(crossrefs_dir.glob("*.json"))

        if not map_files:
            return {
                "success": True,
                "has_evidence_map": False,
                "status": "not_started",
                "status_color": "secondary",
                "message": "No evidence mapping has been performed",
                "recommended_actions": ["Run evidence mapping workflow"]
            }

        # Find latest by modification time
        latest_file = max(map_files, key=lambda p: p.stat().st_mtime)

        try:
            with open(latest_file, "r", encoding="utf-8") as f:
                evidence_map = json.load(f)
        except Exception as e:
            return {"error": f"Failed to load evidence map: {str(e)}"}

        # Get statistics from map
        stats = evidence_map.get("statistics", evidence_map.get("coverage_stats", {}))

        # Calculate coverage if not in stats
        if "coverage_percent" not in stats:
            entries = evidence_map.get("entries", evidence_map.get("mappings", []))
            total = len(entries)
            with_evidence = sum(
                1 for e in entries
                if e.get("evidence_found") or e.get("status") in ["satisfied", "partial"]
            )
            coverage = round(with_evidence / total * 100, 1) if total > 0 else 0
            stats["coverage_percent"] = coverage
            stats["total_standards"] = total
            stats["with_evidence"] = with_evidence

        coverage = stats.get("coverage_percent", 0)

        # Determine overall status
        if coverage >= 90:
            status = "ready"
            status_color = "success"
        elif coverage >= 70:
            status = "mostly_ready"
            status_color = "warning"
        elif coverage >= 50:
            status = "significant_gaps"
            status_color = "warning"
        else:
            status = "not_ready"
            status_color = "danger"

        return {
            "success": True,
            "has_evidence_map": True,
            "evidence_map_id": evidence_map.get("id", latest_file.stem),
            "standards_library_id": evidence_map.get("standards_library_id", standards_id),
            "status": status,
            "status_color": status_color,
            "coverage_stats": {
                "total": stats.get("total_standards", 0),
                "with_evidence": stats.get("with_evidence", 0),
                "strong": stats.get("strong", 0),
                "adequate": stats.get("adequate", 0),
                "weak": stats.get("weak", 0),
                "missing": stats.get("missing", 0),
                "coverage_percent": coverage
            },
            "last_updated": evidence_map.get("created_at", ""),
            "source_file": str(latest_file),
            "recommended_actions": self._get_recommended_actions(stats)
        }

    def _get_recommended_actions(self, stats: Dict[str, Any]) -> List[str]:
        """Generate recommended actions based on coverage stats."""
        actions = []
        missing = stats.get("missing", 0)
        weak = stats.get("weak", 0)
        coverage = stats.get("coverage_percent", 0)

        if missing > 0:
            actions.append(f"Upload documents for {missing} standards with no evidence")
        if weak > 0:
            actions.append(f"Improve evidence quality for {weak} weak mappings")
        if coverage < 80:
            actions.append("Run gap analysis to identify priority areas")
        if coverage >= 80 and coverage < 100:
            actions.append("Review partial mappings to strengthen coverage")
        if coverage >= 90:
            actions.append("Ready for crosswalk export and submission preparation")

        return actions if actions else ["Evidence mapping is complete"]

    def run_workflow(self, action: str, inputs: Dict[str, Any]) -> AgentResult:
        """Run an evidence mapping workflow."""
        if action == "map_all_standards":
            return self._workflow_map_all_standards(inputs)
        elif action == "gap_analysis":
            return self._workflow_gap_analysis(inputs)

        return AgentResult.error(f"Unknown workflow action: {action}")

    def _workflow_map_all_standards(self, inputs: Dict[str, Any]) -> AgentResult:
        """Map all standards to evidence for an institution."""
        from src.core.standards_store import get_standards_store
        institution_id = inputs.get("institution_id")
        standards_id = inputs.get("standards_id")
        save_to_db = inputs.get("save_to_db", False)

        if not institution_id or not standards_id:
            return AgentResult.error("institution_id and standards_id are required")

        # Load standards
        store = get_standards_store()
        library = store.load(standards_id)
        if not library:
            return AgentResult.error(f"Standards library not found: {standards_id}")

        # Build mappings for each checklist item
        mappings = []
        for item in library.checklist_items:
            mapping_result = self._tool_map_standard_to_evidence({
                "standard_id": item.number,
                "standard_text": item.description,
                "institution_id": institution_id
            })

            if "error" not in mapping_result:
                mapping_data = mapping_result.get("mapping", {})
                mappings.append(mapping_data)

        # Calculate coverage statistics
        total = len(mappings)
        satisfied = sum(1 for m in mappings if m.get("status") == "satisfied")
        partial = sum(1 for m in mappings if m.get("status") == "partial")
        weak = sum(1 for m in mappings if m.get("status") == "weak")
        missing = sum(1 for m in mappings if m.get("status") == "missing")

        coverage_stats = {
            "total_standards": total,
            "satisfied": satisfied,
            "partial": partial,
            "weak": weak,
            "missing": missing,
            "with_evidence": satisfied + partial,
            "coverage_percent": round((satisfied + partial) / total * 100, 1) if total > 0 else 0,
            "confidence_avg": round(
                sum(m.get("confidence", 0) for m in mappings) / total, 2
            ) if total > 0 else 0
        }

        # Build evidence map
        evidence_map = EvidenceMap(
            institution_id=institution_id,
            standards_library_id=standards_id,
            mappings=[EvidenceMapping.from_dict(m) for m in mappings],
            coverage_stats=coverage_stats
        )

        # Save to workspace
        save_result = self._tool_save_evidence_map({
            "institution_id": institution_id,
            "evidence_map": evidence_map.to_dict(),
            "filename": f"evidence_map_{standards_id}.json"
        })

        # Determine if human review needed
        needs_review = coverage_stats["coverage_percent"] < 70 or missing > 0

        # Build next actions
        next_actions = []
        if missing > 0:
            next_actions.append({
                "action": "identify_evidence_gaps",
                "priority": "high",
                "reason": f"{missing} standards lack evidence"
            })
        next_actions.append({
            "action": "generate_crosswalk_table",
            "priority": "medium",
            "reason": "Export crosswalk for review"
        })

        return AgentResult(
            status="success",
            confidence=coverage_stats["confidence_avg"],
            data={
                "evidence_map_id": evidence_map.id,
                "coverage_stats": coverage_stats,
                "output_path": save_result.get("path"),
                "standards_library": library.name
            },
            artifacts=[save_result.get("path")] if save_result.get("path") else [],
            human_checkpoint_required=needs_review,
            checkpoint_reason=f"Coverage {coverage_stats['coverage_percent']}% - {missing} gaps need attention" if needs_review else "",
            next_actions=next_actions
        )

    def _workflow_gap_analysis(self, inputs: Dict[str, Any]) -> AgentResult:
        """Perform gap analysis on evidence mapping."""
        institution_id = inputs.get("institution_id")
        standards_id = inputs.get("standards_id")
        evidence_map_id = inputs.get("evidence_map_id")

        if not institution_id:
            return AgentResult.error("institution_id is required")

        # If no standards_id provided, try to infer from existing map
        if not standards_id and evidence_map_id:
            map_data = self._load_evidence_map(institution_id, evidence_map_id)
            if map_data:
                standards_id = map_data.get("standards_library_id")

        if not standards_id:
            return AgentResult.error("standards_id is required (could not infer from evidence map)")

        # Run tool-based gap analysis
        gaps_result = self._tool_identify_evidence_gaps({
            "institution_id": institution_id,
            "standards_id": standards_id,
            "evidence_map_id": evidence_map_id
        })

        if "error" in gaps_result:
            return AgentResult.error(gaps_result["error"])

        # Cross-check with Evidence Contract Service if available
        contract_coverage = None
        try:
            from src.services.evidence_contract_service import check_evidence_coverage
            accreditor_code = standards_id.replace("std_", "").upper() if standards_id else "ACCSC"
            coverage_report = check_evidence_coverage(
                institution_id,
                accreditor_code=accreditor_code
            )
            contract_coverage = coverage_report.to_dict()
        except Exception:
            # Service may not be available or DB not set up
            pass

        # Build combined report
        critical_count = gaps_result.get("summary", {}).get("critical_count", 0)
        total_gaps = gaps_result.get("total_gaps", 0)

        combined_report = {
            "agent_analysis": gaps_result,
            "contract_coverage": contract_coverage,
            "critical_count": critical_count,
            "total_gaps": total_gaps,
            "recommendations": self._generate_recommendations(gaps_result)
        }

        # Determine confidence - lower with more critical gaps
        confidence = max(0.3, 0.9 - (critical_count * 0.1))

        # Build next actions
        next_actions = []
        if critical_count > 0:
            next_actions.append({
                "action": "remediation",
                "priority": "critical",
                "agent": "REMEDIATION",
                "reason": f"{critical_count} critical gaps require attention"
            })
        if total_gaps > 0:
            next_actions.append({
                "action": "document_upload",
                "priority": "high",
                "reason": "Upload additional evidence documents"
            })

        return AgentResult(
            status="success",
            confidence=confidence,
            data=combined_report,
            human_checkpoint_required=critical_count > 0,
            checkpoint_reason=f"{critical_count} critical gaps require attention" if critical_count > 0 else "",
            next_actions=next_actions
        )

    def _generate_recommendations(self, gaps_result: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate prioritized recommendations from gaps."""
        recommendations = []

        for gap in gaps_result.get("critical_gaps", []):
            recommendations.append({
                "priority": 1,
                "severity": "critical",
                "action": f"Address critical gap: {gap.get('standard_number', gap.get('standard_id', 'Unknown'))}",
                "details": gap.get("suggestions", []),
                "standard_text": gap.get("standard_text", "")[:100]
            })

        for gap in gaps_result.get("high_gaps", []):
            recommendations.append({
                "priority": 2,
                "severity": "high",
                "action": f"Address high priority gap: {gap.get('standard_number', gap.get('standard_id', 'Unknown'))}",
                "details": gap.get("suggestions", []),
                "standard_text": gap.get("standard_text", "")[:100]
            })

        for gap in gaps_result.get("advisory_gaps", [])[:5]:  # Limit advisory
            recommendations.append({
                "priority": 3,
                "severity": "advisory",
                "action": f"Consider strengthening: {gap.get('standard_number', gap.get('standard_id', 'Unknown'))}",
                "details": gap.get("suggestions", []),
                "standard_text": gap.get("standard_text", "")[:100]
            })

        return sorted(recommendations, key=lambda x: x["priority"])
