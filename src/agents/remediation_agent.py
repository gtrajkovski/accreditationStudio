"""Remediation Agent.

Produces redlines, clean finals, and crossrefs from audit findings.
Never modifies originals - all changes are versioned in the workspace.

Tools:
- load_audit_findings: Load findings from a completed audit
- generate_correction: Generate corrected text for a finding
- create_redline_document: Create DOCX with tracked changes
- create_final_document: Create clean DOCX with corrections applied
- apply_truth_index: Apply authoritative values from truth index
- save_remediation: Persist remediation result to workspace
"""

import io
import json
import logging
import re
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional, Generator

logger = logging.getLogger(__name__)

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from src.agents.base_agent import BaseAgent, AgentType
from src.agents.registry import register_agent
from src.core.models import (
    AgentResult,
    AgentSession,
    Audit,
    AuditFinding,
    AuditStatus,
    ComplianceStatus,
    Document,
    FindingSeverity,
    RemediationChange,
    RemediationResult,
    RemediationStatus,
    generate_id,
    now_iso,
)


SYSTEM_PROMPT = """You are the Remediation Agent for AccreditAI.

You generate corrected document versions based on compliance audit findings.

CRITICAL RULES:
1. NEVER modify files in originals/ folder - they are read-only
2. Use truth_index.json as the authoritative source for all institutional values
3. Every inserted compliance statement must cite the governing standard
4. Link every change to the finding that triggered it
5. Generate both redline (tracked changes) and clean final versions

REMEDIATION WORKFLOW:
1. LOAD: Load audit findings that need remediation
2. ANALYZE: For each finding, determine the appropriate correction
3. GENERATE: Create corrected text with standard citations
4. BUILD: Create redline document showing changes
5. FINALIZE: Create clean final document
6. APPLY TRUTH: Apply truth index values for consistency
7. SAVE: Persist all outputs to workspace

CHANGE TYPES:
- INSERT: Add new content where none exists
- REPLACE: Replace existing text with corrected version
- DELETE: Remove non-compliant content

OUTPUT FILES:
- redlines/<doc_id>_<timestamp>.docx - Shows tracked changes with strikethrough/highlight
- finals/<doc_id>_<timestamp>.docx - Clean version with all corrections applied
- crossrefs/<doc_id>_<timestamp>.docx - Version with standard citations as comments

QUALITY STANDARDS:
- Every correction must be traceable to a specific finding
- Citations must use exact standard references (e.g., "ACCSC Section I.A.1")
- Language must match institutional tone and style
- Numerical values must come from truth index"""


@register_agent(AgentType.REMEDIATION)
class RemediationAgent(BaseAgent):
    """Remediation Agent.

    Rules:
    - Never modify originals (read-only)
    - Use truth index as authoritative source
    - Every inserted clause linked to finding and standard citation
    - Generate clean final + redline version

    Outputs:
    - redlines/<doc_id>_<timestamp>.docx
    - finals/<doc_id>_<timestamp>.docx
    - crossrefs/<doc_id>_<timestamp>.docx
    """

    def __init__(
        self,
        session: AgentSession,
        workspace_manager=None,
        on_update=None,
    ):
        """Initialize the Remediation Agent."""
        super().__init__(session, workspace_manager, on_update)
        self._institution_id: Optional[str] = None
        self._remediation_cache: Dict[str, RemediationResult] = {}
        self._audit_cache: Dict[str, Audit] = {}

    @property
    def agent_type(self) -> AgentType:
        return AgentType.REMEDIATION

    @property
    def system_prompt(self) -> str:
        return SYSTEM_PROMPT

    @property
    def tools(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "load_audit_findings",
                "description": "Load findings from a completed audit that need remediation. Returns findings with status non_compliant or partial.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "audit_id": {
                            "type": "string",
                            "description": "Audit ID to load findings from"
                        },
                        "severity_filter": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Filter by severity (critical, significant, advisory)"
                        }
                    },
                    "required": ["institution_id", "audit_id"]
                }
            },
            {
                "name": "generate_correction",
                "description": "Generate corrected text for a specific finding using AI.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "finding_id": {
                            "type": "string",
                            "description": "ID of the finding to correct"
                        },
                        "audit_id": {
                            "type": "string",
                            "description": "Audit ID containing the finding"
                        },
                        "context": {
                            "type": "string",
                            "description": "Additional context about the document or institution"
                        }
                    },
                    "required": ["finding_id", "audit_id"]
                }
            },
            {
                "name": "generate_all_corrections",
                "description": "Generate corrections for all non-compliant/partial findings in an audit.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "audit_id": {
                            "type": "string",
                            "description": "Audit ID"
                        },
                        "max_findings": {
                            "type": "integer",
                            "description": "Maximum findings to process (default 20)"
                        }
                    },
                    "required": ["institution_id", "audit_id"]
                }
            },
            {
                "name": "create_redline_document",
                "description": "Create a DOCX document showing changes with tracked changes formatting.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "remediation_id": {
                            "type": "string",
                            "description": "Remediation result ID"
                        }
                    },
                    "required": ["institution_id", "remediation_id"]
                }
            },
            {
                "name": "create_final_document",
                "description": "Create a clean DOCX document with all corrections applied.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "remediation_id": {
                            "type": "string",
                            "description": "Remediation result ID"
                        }
                    },
                    "required": ["institution_id", "remediation_id"]
                }
            },
            {
                "name": "apply_truth_index",
                "description": "Apply truth index values to remediation changes for consistency.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "remediation_id": {
                            "type": "string",
                            "description": "Remediation result ID"
                        },
                        "fields": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Specific fields to apply (optional, applies all if not specified)"
                        }
                    },
                    "required": ["institution_id", "remediation_id"]
                }
            },
            {
                "name": "save_remediation",
                "description": "Save the remediation result and generated documents to workspace.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {
                            "type": "string",
                            "description": "Institution ID"
                        },
                        "remediation_id": {
                            "type": "string",
                            "description": "Remediation result ID"
                        }
                    },
                    "required": ["institution_id", "remediation_id"]
                }
            }
        ]

    def _execute_tool(self, tool_name: str, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Execute a remediation tool."""
        if tool_name == "load_audit_findings":
            return self._tool_load_audit_findings(tool_input)
        elif tool_name == "generate_correction":
            return self._tool_generate_correction(tool_input)
        elif tool_name == "generate_all_corrections":
            return self._tool_generate_all_corrections(tool_input)
        elif tool_name == "create_redline_document":
            return self._tool_create_redline_document(tool_input)
        elif tool_name == "create_final_document":
            return self._tool_create_final_document(tool_input)
        elif tool_name == "apply_truth_index":
            return self._tool_apply_truth_index(tool_input)
        elif tool_name == "save_remediation":
            return self._tool_save_remediation(tool_input)
        return {"error": f"Unknown tool: {tool_name}"}

    # =========================================================================
    # Helper Methods
    # =========================================================================

    def _load_audit(self, institution_id: str, audit_id: str) -> Optional[Audit]:
        """Load an audit from cache or workspace."""
        cache_key = f"{institution_id}:{audit_id}"
        if cache_key in self._audit_cache:
            return self._audit_cache[cache_key]

        if self.workspace_manager:
            audit_path = f"audits/{audit_id}.json"
            try:
                data = self.workspace_manager.read_file(institution_id, audit_path)
                if data:
                    audit = Audit.from_dict(json.loads(data.decode("utf-8")))
                    self._audit_cache[cache_key] = audit
                    return audit
            except Exception as e:
                logger.debug("Failed to load audit %s: %s", audit_id, e)
        return None

    def _load_document(self, institution_id: str, document_id: str) -> Optional[Document]:
        """Load a document from the workspace."""
        if not self.workspace_manager:
            return None
        institution = self.workspace_manager.load_institution(institution_id)
        if not institution:
            return None
        for doc in institution.documents:
            if doc.id == document_id:
                return doc
        return None

    def _get_document_text(self, institution_id: str, document_id: str) -> Optional[str]:
        """Get full extracted text for a document."""
        doc = self._load_document(institution_id, document_id)
        if not doc:
            return None
        if doc.extracted_text:
            return doc.extracted_text
        return None

    def _get_truth_index(self, institution_id: str) -> Optional[Dict[str, Any]]:
        """Load truth index for institution."""
        if not self.workspace_manager:
            return None
        return self.workspace_manager.get_truth_index(institution_id)

    def _get_or_create_remediation(
        self,
        institution_id: str,
        audit_id: str,
        document_id: str
    ) -> RemediationResult:
        """Get existing or create new remediation result."""
        cache_key = f"{institution_id}:{audit_id}"
        if cache_key in self._remediation_cache:
            return self._remediation_cache[cache_key]

        # Try to load from workspace
        if self.workspace_manager:
            remed_path = f"remediations/{audit_id}_remediation.json"
            try:
                data = self.workspace_manager.read_file(institution_id, remed_path)
                if data:
                    remed = RemediationResult.from_dict(json.loads(data.decode("utf-8")))
                    self._remediation_cache[cache_key] = remed
                    return remed
            except Exception as e:
                logger.debug("Failed to load remediation for audit %s: %s", audit_id, e)

        # Create new
        remed = RemediationResult(
            audit_id=audit_id,
            document_id=document_id,
            institution_id=institution_id,
            status=RemediationStatus.PENDING,
            ai_model_used=self.model,
        )
        self._remediation_cache[cache_key] = remed
        return remed

    def _save_remediation_to_workspace(self, remediation: RemediationResult) -> str:
        """Persist remediation result to workspace."""
        if not self.workspace_manager:
            return ""

        remed_path = f"remediations/{remediation.audit_id}_remediation.json"
        self.workspace_manager.save_file(
            remediation.institution_id,
            remed_path,
            json.dumps(remediation.to_dict(), indent=2).encode("utf-8"),
            create_version=True
        )
        return remed_path

    # =========================================================================
    # Tool Implementations
    # =========================================================================

    def _tool_load_audit_findings(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Load findings from a completed audit."""
        institution_id = tool_input.get("institution_id", "")
        audit_id = tool_input.get("audit_id", "")
        severity_filter = tool_input.get("severity_filter", [])

        if not institution_id or not audit_id:
            return {"error": "institution_id and audit_id are required"}

        self._institution_id = institution_id

        audit = self._load_audit(institution_id, audit_id)
        if not audit:
            return {"error": f"Audit not found: {audit_id}"}

        if audit.status not in [AuditStatus.COMPLETED, AuditStatus.REVIEWED, AuditStatus.APPROVED]:
            return {"error": f"Audit not completed. Status: {audit.status.value}"}

        # Filter findings that need remediation
        findings_to_fix = []
        for finding in audit.findings:
            if finding.status not in [ComplianceStatus.NON_COMPLIANT, ComplianceStatus.PARTIAL]:
                continue
            if severity_filter:
                if finding.severity.value not in severity_filter:
                    continue
            findings_to_fix.append({
                "id": finding.id,
                "item_number": finding.item_number,
                "item_description": finding.item_description,
                "status": finding.status.value,
                "severity": finding.severity.value,
                "finding_detail": finding.finding_detail,
                "recommendation": finding.recommendation,
                "evidence_in_document": finding.evidence_in_document[:500] if finding.evidence_in_document else "",
                "page_numbers": finding.page_numbers,
            })

        # Initialize remediation result
        remed = self._get_or_create_remediation(institution_id, audit_id, audit.document_id)
        remed.status = RemediationStatus.IN_PROGRESS

        return {
            "success": True,
            "audit_id": audit_id,
            "document_id": audit.document_id,
            "remediation_id": remed.id,
            "total_findings": len(audit.findings),
            "findings_needing_remediation": len(findings_to_fix),
            "by_severity": {
                "critical": len([f for f in findings_to_fix if f["severity"] == "critical"]),
                "significant": len([f for f in findings_to_fix if f["severity"] == "significant"]),
                "advisory": len([f for f in findings_to_fix if f["severity"] == "advisory"]),
            },
            "findings": findings_to_fix,
            "message": f"Loaded {len(findings_to_fix)} findings requiring remediation"
        }

    def _tool_generate_correction(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Generate corrected text for a specific finding."""
        finding_id = tool_input.get("finding_id", "")
        audit_id = tool_input.get("audit_id", "")
        context = tool_input.get("context", "")

        if not finding_id or not audit_id:
            return {"error": "finding_id and audit_id are required"}

        # Find the audit and finding
        audit = None
        for cache_key, cached_audit in self._audit_cache.items():
            if cached_audit.id == audit_id:
                audit = cached_audit
                break

        if not audit:
            return {"error": f"Audit not loaded. Call load_audit_findings first."}

        # Find the specific finding
        finding = None
        for f in audit.findings:
            if f.id == finding_id:
                finding = f
                break

        if not finding:
            return {"error": f"Finding not found: {finding_id}"}

        # Generate correction using AI
        prompt = f"""Generate corrected compliance text for this finding:

FINDING:
- Item: {finding.item_number}
- Description: {finding.item_description}
- Status: {finding.status.value}
- Severity: {finding.severity.value}
- Issue: {finding.finding_detail}
- Current Text: {finding.evidence_in_document[:800] if finding.evidence_in_document else 'No current text'}
- Recommendation: {finding.recommendation}

{f"Additional Context: {context}" if context else ""}

Generate a correction that:
1. Fully addresses the compliance requirement
2. Uses clear, professional language
3. Cites the governing standard where appropriate
4. Can be inserted into an institutional document

Respond with JSON in this format:
{{"change_type": "insert|replace", "corrected_text": "The corrected text...", "rationale": "Why this correction addresses the finding", "standard_citation": "e.g., ACCSC Section I.A.1", "confidence": 0.0-1.0}}"""

        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=1024,
                system="You are an expert compliance writer. Generate professional, accurate corrections for accreditation documents. Output only valid JSON.",
                messages=[{"role": "user", "content": prompt}]
            )

            response_text = response.content[0].text.strip()

            # Parse JSON from response
            if response_text.startswith("{"):
                correction_data = json.loads(response_text)
            else:
                json_match = re.search(r'\{[^{}]*\}', response_text, re.DOTALL)
                if json_match:
                    correction_data = json.loads(json_match.group())
                else:
                    correction_data = {
                        "change_type": "insert",
                        "corrected_text": response_text[:500],
                        "rationale": "Generated from AI response",
                        "standard_citation": finding.item_number,
                        "confidence": 0.5
                    }

            # Create remediation change
            change = RemediationChange(
                finding_id=finding_id,
                item_number=finding.item_number,
                change_type=correction_data.get("change_type", "insert"),
                location=f"Page {finding.page_numbers}" if finding.page_numbers else "TBD",
                original_text=finding.evidence_in_document[:500] if finding.evidence_in_document else "",
                corrected_text=correction_data.get("corrected_text", ""),
                standard_citation=correction_data.get("standard_citation", finding.item_number),
                rationale=correction_data.get("rationale", ""),
                ai_confidence=correction_data.get("confidence", 0.7),
            )

            # Add to remediation result
            cache_key = f"{self._institution_id}:{audit_id}"
            if cache_key in self._remediation_cache:
                remed = self._remediation_cache[cache_key]
                # Check if change already exists
                existing = [c for c in remed.changes if c.finding_id == finding_id]
                if existing:
                    remed.changes.remove(existing[0])
                remed.changes.append(change)
                remed.findings_addressed = len(remed.changes)

            return {
                "success": True,
                "change": change.to_dict(),
                "message": f"Generated correction for {finding.item_number}"
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to generate correction: {str(e)}",
                "finding_id": finding_id
            }

    def _tool_generate_all_corrections(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Generate corrections for all findings needing remediation."""
        institution_id = tool_input.get("institution_id", "")
        audit_id = tool_input.get("audit_id", "")
        max_findings = tool_input.get("max_findings", 20)

        if not institution_id or not audit_id:
            return {"error": "institution_id and audit_id are required"}

        self._institution_id = institution_id

        # Load audit
        audit = self._load_audit(institution_id, audit_id)
        if not audit:
            return {"error": f"Audit not found: {audit_id}"}

        # Get remediation result
        remed = self._get_or_create_remediation(institution_id, audit_id, audit.document_id)

        # Filter findings needing remediation
        findings_to_fix = [
            f for f in audit.findings
            if f.status in [ComplianceStatus.NON_COMPLIANT, ComplianceStatus.PARTIAL]
        ][:max_findings]

        # Load truth index for context
        truth_index = self._get_truth_index(institution_id)
        context = ""
        if truth_index:
            inst_info = truth_index.get("institution", {})
            context = f"Institution: {inst_info.get('name', 'Unknown')}"

        generated = []
        skipped = []

        for finding in findings_to_fix:
            result = self._tool_generate_correction({
                "finding_id": finding.id,
                "audit_id": audit_id,
                "context": context
            })

            if result.get("success"):
                generated.append({
                    "finding_id": finding.id,
                    "item_number": finding.item_number,
                    "change_type": result["change"].get("change_type"),
                    "confidence": result["change"].get("ai_confidence"),
                })
            else:
                skipped.append({
                    "finding_id": finding.id,
                    "item_number": finding.item_number,
                    "error": result.get("error", "Unknown error")
                })

        remed.findings_addressed = len(generated)
        remed.findings_skipped = len(skipped)
        remed.status = RemediationStatus.GENERATED

        return {
            "success": True,
            "remediation_id": remed.id,
            "findings_processed": len(findings_to_fix),
            "corrections_generated": len(generated),
            "corrections_skipped": len(skipped),
            "generated": generated,
            "skipped": skipped,
            "message": f"Generated {len(generated)} corrections, skipped {len(skipped)}"
        }

    def _tool_create_redline_document(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Create a DOCX document with redline (tracked changes) formatting."""
        institution_id = tool_input.get("institution_id", "")
        remediation_id = tool_input.get("remediation_id", "")

        if not institution_id or not remediation_id:
            return {"error": "institution_id and remediation_id are required"}

        # Find remediation result
        remed = None
        for cache_key, cached_remed in self._remediation_cache.items():
            if cached_remed.id == remediation_id:
                remed = cached_remed
                break

        if not remed:
            return {"error": f"Remediation not found: {remediation_id}. Run generate_all_corrections first."}

        if not remed.changes:
            return {"error": "No changes to create redline from"}

        # Create DOCX document
        doc = DocxDocument()

        # Title
        title = doc.add_heading("Compliance Remediation - Redline Version", 0)

        # Metadata section
        doc.add_paragraph(f"Document ID: {remed.document_id}")
        doc.add_paragraph(f"Audit ID: {remed.audit_id}")
        doc.add_paragraph(f"Generated: {now_iso()}")
        doc.add_paragraph(f"Changes: {len(remed.changes)}")
        doc.add_paragraph()

        doc.add_heading("Changes Required", level=1)

        # Add each change with redline formatting
        for idx, change in enumerate(remed.changes, 1):
            # Change header
            doc.add_heading(f"Change {idx}: {change.item_number}", level=2)

            # Citation
            citation_para = doc.add_paragraph()
            citation_para.add_run("Standard: ").bold = True
            citation_para.add_run(change.standard_citation)

            # Location
            if change.location:
                loc_para = doc.add_paragraph()
                loc_para.add_run("Location: ").bold = True
                loc_para.add_run(change.location)

            # Original text (strikethrough style)
            if change.original_text:
                doc.add_paragraph()
                orig_para = doc.add_paragraph()
                orig_para.add_run("Original (DELETE): ").bold = True
                orig_run = orig_para.add_run(change.original_text[:500])
                orig_run.font.strike = True
                orig_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            # Corrected text (highlight style)
            doc.add_paragraph()
            new_para = doc.add_paragraph()
            new_para.add_run("Corrected (INSERT): ").bold = True
            new_run = new_para.add_run(change.corrected_text)
            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            new_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green

            # Rationale
            if change.rationale:
                doc.add_paragraph()
                rat_para = doc.add_paragraph()
                rat_para.add_run("Rationale: ").bold = True
                rat_para.add_run(change.rationale)

            # Confidence
            conf_para = doc.add_paragraph()
            conf_para.add_run("AI Confidence: ").bold = True
            conf_para.add_run(f"{change.ai_confidence:.0%}")

            doc.add_paragraph()  # Spacing

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Save to workspace
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        redline_filename = f"redlines/{remed.document_id}_{timestamp}_redline.docx"

        if self.workspace_manager:
            self.workspace_manager.save_file(
                institution_id,
                redline_filename,
                doc_bytes.read(),
                create_version=False
            )
            remed.redline_path = redline_filename

        return {
            "success": True,
            "path": redline_filename,
            "changes_included": len(remed.changes),
            "message": f"Created redline document with {len(remed.changes)} changes"
        }

    def _tool_create_final_document(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Create a clean DOCX document with all corrections applied."""
        institution_id = tool_input.get("institution_id", "")
        remediation_id = tool_input.get("remediation_id", "")

        if not institution_id or not remediation_id:
            return {"error": "institution_id and remediation_id are required"}

        # Find remediation result
        remed = None
        for cache_key, cached_remed in self._remediation_cache.items():
            if cached_remed.id == remediation_id:
                remed = cached_remed
                break

        if not remed:
            return {"error": f"Remediation not found: {remediation_id}"}

        if not remed.changes:
            return {"error": "No changes to apply"}

        # Create DOCX document
        doc = DocxDocument()

        # Title
        doc.add_heading("Compliance Remediation - Final Version", 0)

        # Metadata
        doc.add_paragraph(f"Document ID: {remed.document_id}")
        doc.add_paragraph(f"Audit ID: {remed.audit_id}")
        doc.add_paragraph(f"Generated: {now_iso()}")
        doc.add_paragraph(f"Corrections Applied: {len(remed.changes)}")
        doc.add_paragraph()

        doc.add_heading("Corrected Content", level=1)

        # Group changes by section/category
        changes_by_item = {}
        for change in remed.changes:
            item_num = change.item_number.split(".")[0] if "." in change.item_number else change.item_number
            if item_num not in changes_by_item:
                changes_by_item[item_num] = []
            changes_by_item[item_num].append(change)

        # Add corrected content organized by section
        for section, changes in changes_by_item.items():
            doc.add_heading(f"Section: {section}", level=2)

            for change in changes:
                # Standard reference
                ref_para = doc.add_paragraph()
                ref_para.add_run(f"[{change.standard_citation}] ").italic = True

                # Corrected text
                doc.add_paragraph(change.corrected_text)
                doc.add_paragraph()  # Spacing

        # Summary table
        doc.add_heading("Remediation Summary", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item"
        hdr_cells[1].text = "Change Type"
        hdr_cells[2].text = "Citation"
        hdr_cells[3].text = "Confidence"

        for change in remed.changes:
            row_cells = table.add_row().cells
            row_cells[0].text = change.item_number
            row_cells[1].text = change.change_type.upper()
            row_cells[2].text = change.standard_citation
            row_cells[3].text = f"{change.ai_confidence:.0%}"

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Save to workspace
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        final_filename = f"finals/{remed.document_id}_{timestamp}_final.docx"

        if self.workspace_manager:
            self.workspace_manager.save_file(
                institution_id,
                final_filename,
                doc_bytes.read(),
                create_version=False
            )
            remed.final_path = final_filename

        return {
            "success": True,
            "path": final_filename,
            "corrections_applied": len(remed.changes),
            "message": f"Created final document with {len(remed.changes)} corrections applied"
        }

    def _tool_apply_truth_index(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Apply truth index values to remediation changes."""
        institution_id = tool_input.get("institution_id", "")
        remediation_id = tool_input.get("remediation_id", "")
        fields = tool_input.get("fields", [])

        if not institution_id or not remediation_id:
            return {"error": "institution_id and remediation_id are required"}

        # Find remediation result
        remed = None
        for cache_key, cached_remed in self._remediation_cache.items():
            if cached_remed.id == remediation_id:
                remed = cached_remed
                break

        if not remed:
            return {"error": f"Remediation not found: {remediation_id}"}

        # Load truth index
        truth_index = self._get_truth_index(institution_id)
        if not truth_index:
            return {"error": "Truth index not found for institution"}

        # Extract values to apply
        inst_values = truth_index.get("institution", {})
        program_values = truth_index.get("programs", {})
        policy_values = truth_index.get("policies", {})

        # Build replacement map
        replacements = {}

        # Institution values
        if inst_values.get("name"):
            replacements["[INSTITUTION_NAME]"] = inst_values["name"]
            replacements["[institution name]"] = inst_values["name"]

        # Program values (if specific program)
        for prog_id, prog_data in program_values.items():
            if prog_data.get("name_en"):
                replacements[f"[PROGRAM_NAME:{prog_id}]"] = prog_data["name_en"]
            if prog_data.get("total_cost"):
                replacements[f"[PROGRAM_COST:{prog_id}]"] = f"${prog_data['total_cost']:,.2f}"
            if prog_data.get("duration_months"):
                replacements[f"[PROGRAM_DURATION:{prog_id}]"] = f"{prog_data['duration_months']} months"

        # Apply replacements to changes
        changes_modified = 0
        truth_changes = []

        for change in remed.changes:
            original_text = change.corrected_text
            modified_text = original_text

            for placeholder, value in replacements.items():
                if placeholder.lower() in modified_text.lower():
                    # Case-insensitive replacement
                    pattern = re.compile(re.escape(placeholder), re.IGNORECASE)
                    modified_text = pattern.sub(value, modified_text)

            if modified_text != original_text:
                change.corrected_text = modified_text
                changes_modified += 1
                truth_changes.append({
                    "change_id": change.id,
                    "item_number": change.item_number,
                    "replacements_applied": [k for k in replacements if k.lower() in original_text.lower()]
                })

        remed.truth_index_applied = True
        remed.truth_index_changes = truth_changes

        return {
            "success": True,
            "changes_modified": changes_modified,
            "replacements_available": list(replacements.keys()),
            "truth_changes": truth_changes,
            "message": f"Applied truth index to {changes_modified} changes"
        }

    def _tool_save_remediation(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Save remediation result to workspace."""
        institution_id = tool_input.get("institution_id", "")
        remediation_id = tool_input.get("remediation_id", "")

        if not institution_id or not remediation_id:
            return {"error": "institution_id and remediation_id are required"}

        # Find remediation result
        remed = None
        for cache_key, cached_remed in self._remediation_cache.items():
            if cached_remed.id == remediation_id:
                remed = cached_remed
                break

        if not remed:
            return {"error": f"Remediation not found: {remediation_id}"}

        # Update status
        remed.completed_at = now_iso()
        if remed.redline_path and remed.final_path:
            remed.status = RemediationStatus.GENERATED

        # Save to workspace
        saved_path = self._save_remediation_to_workspace(remed)

        # Track artifact
        self.session.artifacts_created.append(saved_path)
        if remed.redline_path:
            self.session.artifacts_created.append(remed.redline_path)
        if remed.final_path:
            self.session.artifacts_created.append(remed.final_path)

        return {
            "success": True,
            "remediation_id": remed.id,
            "status": remed.status.value,
            "saved_path": saved_path,
            "artifacts": {
                "remediation_json": saved_path,
                "redline_docx": remed.redline_path,
                "final_docx": remed.final_path,
            },
            "summary": {
                "findings_addressed": remed.findings_addressed,
                "findings_skipped": remed.findings_skipped,
                "changes_count": len(remed.changes),
                "truth_index_applied": remed.truth_index_applied,
            },
            "message": f"Saved remediation with {len(remed.changes)} changes"
        }

    # =========================================================================
    # Workflow Methods
    # =========================================================================

    def remediate_document(
        self,
        institution_id: str,
        audit_id: str,
        max_findings: int = 20,
    ) -> Generator[Dict[str, Any], None, Dict[str, Any]]:
        """High-level entry point for full document remediation.

        Args:
            institution_id: Institution ID
            audit_id: Completed audit ID
            max_findings: Maximum findings to process

        Yields:
            Progress updates during remediation

        Returns:
            Final remediation result
        """
        self._institution_id = institution_id

        prompt = f"""Remediate the document based on audit findings:

Institution: {institution_id}
Audit ID: {audit_id}
Max Findings: {max_findings}

Execute the full remediation workflow:
1. load_audit_findings - Load findings needing remediation
2. generate_all_corrections - Generate corrections for all findings
3. apply_truth_index - Apply institutional values from truth index
4. create_redline_document - Create redline version with tracked changes
5. create_final_document - Create clean final version
6. save_remediation - Save all results to workspace

Report your progress after each step."""

        for update in self.run_turn(prompt):
            yield update

        return {
            "status": "completed",
            "institution_id": institution_id,
            "audit_id": audit_id,
        }

    def run_programmatic_remediation(
        self,
        institution_id: str,
        audit_id: str,
        max_findings: int = 20,
    ) -> RemediationResult:
        """Run remediation programmatically without AI orchestration."""
        self._institution_id = institution_id

        # Step 1: Load findings
        load_result = self._tool_load_audit_findings({
            "institution_id": institution_id,
            "audit_id": audit_id,
        })

        if "error" in load_result:
            raise ValueError(load_result["error"])

        remediation_id = load_result["remediation_id"]

        # Step 2: Generate all corrections
        gen_result = self._tool_generate_all_corrections({
            "institution_id": institution_id,
            "audit_id": audit_id,
            "max_findings": max_findings,
        })

        if "error" in gen_result:
            raise ValueError(gen_result["error"])

        # Step 3: Apply truth index
        self._tool_apply_truth_index({
            "institution_id": institution_id,
            "remediation_id": remediation_id,
        })

        # Step 4: Create redline
        self._tool_create_redline_document({
            "institution_id": institution_id,
            "remediation_id": remediation_id,
        })

        # Step 5: Create final
        self._tool_create_final_document({
            "institution_id": institution_id,
            "remediation_id": remediation_id,
        })

        # Step 6: Save
        self._tool_save_remediation({
            "institution_id": institution_id,
            "remediation_id": remediation_id,
        })

        # Return the remediation result
        for cache_key, remed in self._remediation_cache.items():
            if remed.id == remediation_id:
                return remed

        raise ValueError("Remediation result not found after processing")

    def run_workflow(self, action: str, inputs: Dict[str, Any]) -> AgentResult:
        """Run a remediation workflow."""
        if action == "full_remediation":
            try:
                remed = self.run_programmatic_remediation(
                    institution_id=inputs.get("institution_id", ""),
                    audit_id=inputs.get("audit_id", ""),
                    max_findings=inputs.get("max_findings", 20),
                )
                return AgentResult.success(
                    data=remed.to_dict(),
                    confidence=0.8,
                    artifacts=[remed.redline_path, remed.final_path],
                )
            except Exception as e:
                return AgentResult.error(str(e))

        return AgentResult.error(f"Unknown workflow: {action}")
