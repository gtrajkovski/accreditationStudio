"""Packet Agent.

Assembles submission-ready response packages and exhibit binders for
accreditation submissions. Validates evidence coverage and generates
professionally formatted DOCX exports.

Tools:
- create_packet: Initialize a new submission packet
- load_findings: Load findings to address in the packet
- add_narrative_section: Add a narrative response section
- add_exhibit: Add an exhibit/evidence document
- build_cover_page: Generate the cover page
- build_table_of_contents: Generate TOC
- build_evidence_index: Create standard → evidence crosswalk
- validate_packet: Check evidence coverage requirements
- export_docx: Export packet as DOCX
- export_zip: Export complete submission as ZIP folder
- save_packet: Persist packet to workspace
"""

import io
import json
import os
import zipfile
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.agents.base_agent import BaseAgent, AgentType
from src.agents.registry import register_agent
from src.core.models import (
    AgentResult,
    AgentSession,
    ExhibitEntry,
    PacketSection,
    PacketSectionType,
    PacketStatus,
    SubmissionPacket,
    SubmissionType,
    ValidationIssue,
    generate_id,
    now_iso,
)
from src.config import Config


SYSTEM_PROMPT = """You are the Packet Agent for AccreditAI.

You assemble submission-ready response packages and exhibit binders for accreditation submissions.

CRITICAL RULES:
1. Every standard in the packet must have evidence linked
2. Critical findings must have narrative responses
3. All sections require human review before export
4. Never fabricate evidence references - only link to real documents
5. Validate coverage before allowing export

PACKET STRUCTURE:
1. Cover Page - Institution name, accreditor, submission type, date
2. Table of Contents - Auto-generated from sections
3. Executive Summary - Optional overview
4. Narrative Responses - Issue-by-issue responses to findings
5. Corrective Action Plan - Table of planned actions
6. Evidence Index - Standard → evidence crosswalk
7. Exhibit List - Organized list of supporting documents
8. Signature/Certification Page

VALIDATION RULES:
- Every addressed standard must have >=1 evidence reference
- Critical findings must have narrative responses
- Low-confidence sections flagged for human review
- Export blocked until validation passes (unless override approved)

OUTPUT FORMATS:
- DOCX: Professional document with proper formatting
- ZIP: Folder structure with all exhibits organized"""


@register_agent(AgentType.PACKET)
class PacketAgent(BaseAgent):
    """Packet Agent for assembling accreditation submissions.

    Creates professional submission packages with:
    - Cover pages and TOC
    - Narrative responses to findings
    - Evidence indexes and exhibit lists
    - Validation of coverage requirements
    - DOCX and ZIP exports
    """

    def __init__(
        self,
        session: AgentSession,
        workspace_manager=None,
        on_update=None,
    ):
        """Initialize the Packet Agent."""
        super().__init__(session, workspace_manager, on_update)
        self._packet_cache: Dict[str, SubmissionPacket] = {}
        self._institution_id: Optional[str] = None

    @property
    def agent_type(self) -> AgentType:
        return AgentType.PACKET

    @property
    def system_prompt(self) -> str:
        return SYSTEM_PROMPT

    @property
    def tools(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "create_packet",
                "description": "Create a new submission packet.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "institution_id": {"type": "string"},
                        "name": {"type": "string", "description": "Packet name"},
                        "submission_type": {
                            "type": "string",
                            "enum": [
                                "initial_accreditation", "renewal", "substantive_change",
                                "compliance_report", "response_to_findings", "annual_report",
                                "teach_out_plan", "other"
                            ],
                            "default": "response_to_findings"
                        },
                        "accrediting_body": {"type": "string", "description": "e.g., ACCSC, SACSCOC"},
                        "description": {"type": "string"},
                    },
                    "required": ["institution_id", "name", "accrediting_body"],
                },
            },
            {
                "name": "load_findings_for_packet",
                "description": "Load findings from a findings report to address in the packet.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                        "findings_report_id": {"type": "string"},
                        "severity_filter": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Filter by severity (critical, significant, advisory)"
                        },
                    },
                    "required": ["packet_id", "findings_report_id"],
                },
            },
            {
                "name": "add_narrative_section",
                "description": "Add a narrative response section to the packet.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                        "title": {"type": "string"},
                        "content": {"type": "string"},
                        "finding_id": {"type": "string", "description": "Finding being addressed"},
                        "standard_refs": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Standard references (e.g., ['I.A.1', 'I.A.2'])"
                        },
                        "evidence_refs": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Document IDs for evidence"
                        },
                    },
                    "required": ["packet_id", "title", "content"],
                },
            },
            {
                "name": "add_exhibit",
                "description": "Add an exhibit (evidence document) to the packet.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                        "exhibit_number": {"type": "string", "description": "e.g., A-1, B-3"},
                        "title": {"type": "string"},
                        "description": {"type": "string"},
                        "document_id": {"type": "string"},
                        "file_path": {"type": "string"},
                        "standard_refs": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "finding_refs": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                    },
                    "required": ["packet_id", "exhibit_number", "title"],
                },
            },
            {
                "name": "generate_cover_page",
                "description": "Generate the cover page content using AI.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                        "institution_name": {"type": "string"},
                        "submission_date": {"type": "string", "description": "Date in YYYY-MM-DD format"},
                        "contact_name": {"type": "string"},
                        "contact_title": {"type": "string"},
                    },
                    "required": ["packet_id", "institution_name"],
                },
            },
            {
                "name": "build_evidence_index",
                "description": "Build the evidence index (standard → evidence crosswalk).",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                    },
                    "required": ["packet_id"],
                },
            },
            {
                "name": "validate_packet",
                "description": "Validate the packet for evidence coverage requirements.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                        "strict": {
                            "type": "boolean",
                            "default": True,
                            "description": "If true, errors block export"
                        },
                    },
                    "required": ["packet_id"],
                },
            },
            {
                "name": "export_docx",
                "description": "Export the packet as a professional DOCX document.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                        "include_exhibits": {
                            "type": "boolean",
                            "default": False,
                            "description": "Include exhibit content in document"
                        },
                    },
                    "required": ["packet_id"],
                },
            },
            {
                "name": "export_zip",
                "description": "Export complete submission as a ZIP folder with exhibits.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                    },
                    "required": ["packet_id"],
                },
            },
            {
                "name": "get_packet_summary",
                "description": "Get summary statistics for a packet.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                    },
                    "required": ["packet_id"],
                },
            },
            {
                "name": "save_packet",
                "description": "Save the packet to workspace.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "packet_id": {"type": "string"},
                    },
                    "required": ["packet_id"],
                },
            },
        ]

    def _execute_tool(self, tool_name: str, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Execute a tool by name."""
        handlers = {
            "create_packet": self._tool_create_packet,
            "load_findings_for_packet": self._tool_load_findings,
            "add_narrative_section": self._tool_add_narrative,
            "add_exhibit": self._tool_add_exhibit,
            "generate_cover_page": self._tool_generate_cover,
            "build_evidence_index": self._tool_build_evidence_index,
            "validate_packet": self._tool_validate,
            "export_docx": self._tool_export_docx,
            "export_zip": self._tool_export_zip,
            "get_packet_summary": self._tool_get_summary,
            "save_packet": self._tool_save,
        }
        handler = handlers.get(tool_name)
        if handler:
            return handler(tool_input)
        return {"error": f"Unknown tool: {tool_name}"}

    def _get_packet(self, packet_id: str) -> Optional[SubmissionPacket]:
        """Get a packet from cache."""
        return self._packet_cache.get(packet_id)

    def _tool_create_packet(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Create a new submission packet."""
        institution_id = params.get("institution_id", "")
        name = params.get("name", "")
        submission_type = params.get("submission_type", "response_to_findings")
        accrediting_body = params.get("accrediting_body", "")
        description = params.get("description", "")

        if not institution_id or not name or not accrediting_body:
            return {"error": "institution_id, name, and accrediting_body are required"}

        self._institution_id = institution_id

        # Create packet
        packet = SubmissionPacket(
            institution_id=institution_id,
            name=name,
            submission_type=SubmissionType(submission_type),
            accrediting_body=accrediting_body,
            description=description,
            status=PacketStatus.DRAFT,
        )

        self._packet_cache[packet.id] = packet

        return {
            "success": True,
            "packet_id": packet.id,
            "name": name,
            "submission_type": submission_type,
            "accrediting_body": accrediting_body,
            "message": f"Created packet: {name}",
        }

    def _tool_load_findings(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load findings from a findings report."""
        packet_id = params.get("packet_id", "")
        findings_report_id = params.get("findings_report_id", "")
        severity_filter = params.get("severity_filter", [])

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        if not self.workspace_manager:
            return {"error": "Workspace manager not available"}

        # Load findings report
        report_path = f"findings/{findings_report_id}.json"
        report_data = self.workspace_manager.load_file(packet.institution_id, report_path)

        if not report_data:
            return {"error": f"Findings report not found: {findings_report_id}"}

        packet.findings_report_id = findings_report_id

        # Extract findings
        findings = report_data.get("findings", [])
        if severity_filter:
            findings = [f for f in findings if f.get("severity") in severity_filter]

        # Create placeholder sections for each finding
        sections_created = 0
        for idx, finding in enumerate(findings):
            section = PacketSection(
                section_type=PacketSectionType.NARRATIVE_RESPONSE,
                title=f"Response to {finding.get('item_number', 'Finding')}",
                order=idx + 10,  # Leave room for cover/TOC
                finding_id=finding.get("id", ""),
                standard_refs=[finding.get("item_number", "")] if finding.get("item_number") else [],
                content="",  # To be filled by narrative
            )
            packet.sections.append(section)
            sections_created += 1

        packet.update_stats()

        return {
            "success": True,
            "findings_loaded": len(findings),
            "sections_created": sections_created,
            "severity_breakdown": {
                "critical": len([f for f in findings if f.get("severity") == "critical"]),
                "significant": len([f for f in findings if f.get("severity") == "significant"]),
                "advisory": len([f for f in findings if f.get("severity") == "advisory"]),
            },
            "message": f"Loaded {len(findings)} findings, created {sections_created} response sections",
        }

    def _tool_add_narrative(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Add a narrative section to the packet."""
        packet_id = params.get("packet_id", "")
        title = params.get("title", "")
        content = params.get("content", "")
        finding_id = params.get("finding_id", "")
        standard_refs = params.get("standard_refs", [])
        evidence_refs = params.get("evidence_refs", [])

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        # Check if updating existing section
        existing = None
        if finding_id:
            for section in packet.sections:
                if section.finding_id == finding_id:
                    existing = section
                    break

        if existing:
            existing.title = title or existing.title
            existing.content = content
            existing.standard_refs = standard_refs or existing.standard_refs
            existing.evidence_refs = evidence_refs or existing.evidence_refs
            existing.word_count = len(content.split())
            existing.ai_generated = True
            section_id = existing.id
        else:
            section = PacketSection(
                section_type=PacketSectionType.NARRATIVE_RESPONSE,
                title=title,
                content=content,
                finding_id=finding_id,
                standard_refs=standard_refs,
                evidence_refs=evidence_refs,
                word_count=len(content.split()),
                order=len(packet.sections) + 1,
                ai_generated=True,
            )
            packet.sections.append(section)
            section_id = section.id

        packet.update_stats()

        return {
            "success": True,
            "section_id": section_id,
            "title": title,
            "word_count": len(content.split()),
            "updated": existing is not None,
            "message": f"{'Updated' if existing else 'Added'} narrative section: {title}",
        }

    def _tool_add_exhibit(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Add an exhibit to the packet."""
        packet_id = params.get("packet_id", "")
        exhibit_number = params.get("exhibit_number", "")
        title = params.get("title", "")
        description = params.get("description", "")
        document_id = params.get("document_id", "")
        file_path = params.get("file_path", "")
        standard_refs = params.get("standard_refs", [])
        finding_refs = params.get("finding_refs", [])

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        # Determine file type
        file_type = ""
        if file_path:
            file_type = file_path.split(".")[-1].lower() if "." in file_path else ""

        exhibit = ExhibitEntry(
            exhibit_number=exhibit_number,
            title=title,
            description=description,
            document_id=document_id,
            file_path=file_path,
            standard_refs=standard_refs,
            finding_refs=finding_refs,
            file_type=file_type,
        )

        packet.exhibits.append(exhibit)
        packet.update_stats()

        return {
            "success": True,
            "exhibit_id": exhibit.id,
            "exhibit_number": exhibit_number,
            "title": title,
            "total_exhibits": len(packet.exhibits),
            "message": f"Added exhibit {exhibit_number}: {title}",
        }

    def _tool_generate_cover(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Generate the cover page content."""
        packet_id = params.get("packet_id", "")
        institution_name = params.get("institution_name", "")
        submission_date = params.get("submission_date", datetime.now().strftime("%Y-%m-%d"))
        contact_name = params.get("contact_name", "")
        contact_title = params.get("contact_title", "")

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        # Build cover page content
        submission_type_display = packet.submission_type.value.replace("_", " ").title()

        cover_content = f"""
{institution_name.upper()}

{submission_type_display}

Submitted to: {packet.accrediting_body}

Submission Date: {submission_date}

{f"Prepared by: {contact_name}" if contact_name else ""}
{f"{contact_title}" if contact_title else ""}

{packet.description}
""".strip()

        # Add or update cover page section
        cover_section = None
        for section in packet.sections:
            if section.section_type == PacketSectionType.COVER_PAGE:
                cover_section = section
                break

        if cover_section:
            cover_section.content = cover_content
        else:
            cover_section = PacketSection(
                section_type=PacketSectionType.COVER_PAGE,
                title="Cover Page",
                content=cover_content,
                order=0,
            )
            packet.sections.insert(0, cover_section)

        # Re-sort sections by order
        packet.sections.sort(key=lambda s: s.order)
        packet.update_stats()

        return {
            "success": True,
            "section_id": cover_section.id,
            "institution_name": institution_name,
            "submission_type": submission_type_display,
            "message": "Generated cover page",
        }

    def _tool_build_evidence_index(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Build the evidence index crosswalk."""
        packet_id = params.get("packet_id", "")

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        # Build standard → evidence mapping
        evidence_map: Dict[str, List[str]] = {}

        # From sections
        for section in packet.sections:
            for std_ref in section.standard_refs:
                if std_ref not in evidence_map:
                    evidence_map[std_ref] = []
                evidence_map[std_ref].extend(section.evidence_refs)

        # From exhibits
        for exhibit in packet.exhibits:
            for std_ref in exhibit.standard_refs:
                if std_ref not in evidence_map:
                    evidence_map[std_ref] = []
                if exhibit.file_path:
                    evidence_map[std_ref].append(f"Exhibit {exhibit.exhibit_number}: {exhibit.title}")

        # Build content
        index_lines = ["Standard | Evidence Documents", "---|---"]
        for std_ref in sorted(evidence_map.keys()):
            evidence_list = list(set(evidence_map[std_ref]))  # Dedupe
            evidence_str = "; ".join(evidence_list[:5])  # Max 5 per row
            if len(evidence_list) > 5:
                evidence_str += f" (+{len(evidence_list) - 5} more)"
            index_lines.append(f"{std_ref} | {evidence_str}")

        index_content = "\n".join(index_lines)

        # Add or update evidence index section
        index_section = None
        for section in packet.sections:
            if section.section_type == PacketSectionType.EVIDENCE_INDEX:
                index_section = section
                break

        if index_section:
            index_section.content = index_content
        else:
            index_section = PacketSection(
                section_type=PacketSectionType.EVIDENCE_INDEX,
                title="Evidence Index",
                content=index_content,
                order=900,  # Near end
            )
            packet.sections.append(index_section)

        packet.sections.sort(key=lambda s: s.order)
        packet.update_stats()

        return {
            "success": True,
            "standards_indexed": len(evidence_map),
            "total_evidence_refs": sum(len(v) for v in evidence_map.values()),
            "message": f"Built evidence index for {len(evidence_map)} standards",
        }

    def _tool_validate(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Validate the packet for evidence coverage."""
        packet_id = params.get("packet_id", "")
        strict = params.get("strict", True)

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        packet.status = PacketStatus.VALIDATING
        issues: List[ValidationIssue] = []

        # Check 1: All sections need content
        for section in packet.sections:
            if section.section_type == PacketSectionType.NARRATIVE_RESPONSE:
                if not section.content:
                    issues.append(ValidationIssue(
                        issue_type="missing_content",
                        severity="error",
                        message=f"Section '{section.title}' has no content",
                        section_id=section.id,
                        finding_id=section.finding_id,
                    ))

        # Check 2: Standards coverage
        all_standards = set()
        standards_with_evidence = set()

        for section in packet.sections:
            all_standards.update(section.standard_refs)
            if section.evidence_refs:
                standards_with_evidence.update(section.standard_refs)

        for exhibit in packet.exhibits:
            all_standards.update(exhibit.standard_refs)
            if exhibit.file_path or exhibit.document_id:
                standards_with_evidence.update(exhibit.standard_refs)

        missing_evidence = all_standards - standards_with_evidence
        for std_ref in missing_evidence:
            issues.append(ValidationIssue(
                issue_type="missing_evidence",
                severity="warning",
                message=f"Standard {std_ref} has no linked evidence",
                standard_ref=std_ref,
            ))

        # Check 3: Sections without approval
        unapproved = [s for s in packet.sections
                      if s.section_type == PacketSectionType.NARRATIVE_RESPONSE and not s.approved]
        if unapproved:
            issues.append(ValidationIssue(
                issue_type="unapproved_sections",
                severity="warning",
                message=f"{len(unapproved)} narrative section(s) not yet approved",
                can_override=True,
            ))

        # Check 4: Exhibits without files
        missing_files = [e for e in packet.exhibits if e.included and not e.file_path]
        for exhibit in missing_files:
            issues.append(ValidationIssue(
                issue_type="missing_exhibit_file",
                severity="error",
                message=f"Exhibit {exhibit.exhibit_number} has no file path",
            ))

        # Determine validity
        errors = [i for i in issues if i.severity == "error"]
        warnings = [i for i in issues if i.severity == "warning"]

        is_valid = len(errors) == 0
        if strict:
            # In strict mode, non-overridable warnings also fail
            non_overridable_warnings = [w for w in warnings if not w.can_override]
            is_valid = is_valid and len(non_overridable_warnings) == 0

        packet.validation_issues = issues
        packet.is_valid = is_valid
        packet.validated_at = now_iso()
        packet.status = PacketStatus.READY if is_valid else PacketStatus.VALIDATION_FAILED

        return {
            "success": True,
            "is_valid": is_valid,
            "errors": len(errors),
            "warnings": len(warnings),
            "issues": [i.to_dict() for i in issues],
            "can_export": is_valid or not strict,
            "message": "Validation passed" if is_valid else f"Validation failed: {len(errors)} errors, {len(warnings)} warnings",
        }

    def _tool_export_docx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Export packet as DOCX document."""
        packet_id = params.get("packet_id", "")
        include_exhibits = params.get("include_exhibits", False)

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        if not packet.is_valid and packet.status != PacketStatus.APPROVED:
            # Allow export if explicitly approved or validated
            if packet.status not in [PacketStatus.READY, PacketStatus.APPROVED]:
                return {"error": "Packet must pass validation before export. Run validate_packet first."}

        # Create DOCX
        doc = DocxDocument()

        # Sort sections by order
        sections = sorted(packet.sections, key=lambda s: s.order)

        for section in sections:
            if section.section_type == PacketSectionType.COVER_PAGE:
                # Cover page with centered formatting
                for line in section.content.split("\n"):
                    if line.strip():
                        para = doc.add_paragraph(line.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if len(line) < 50:  # Short lines get larger font
                            for run in para.runs:
                                run.font.size = Pt(16)
                                run.font.bold = True
                doc.add_page_break()

            elif section.section_type == PacketSectionType.TABLE_OF_CONTENTS:
                doc.add_heading("Table of Contents", level=1)
                # Auto-generate TOC from sections
                toc_order = 1
                for s in sections:
                    if s.section_type in [PacketSectionType.NARRATIVE_RESPONSE,
                                          PacketSectionType.EVIDENCE_INDEX,
                                          PacketSectionType.EXHIBIT_LIST]:
                        doc.add_paragraph(f"{toc_order}. {s.title}", style="List Number")
                        toc_order += 1
                doc.add_page_break()

            elif section.section_type == PacketSectionType.NARRATIVE_RESPONSE:
                doc.add_heading(section.title, level=2)

                # Standard references
                if section.standard_refs:
                    ref_para = doc.add_paragraph()
                    ref_para.add_run("Standards Addressed: ").bold = True
                    ref_para.add_run(", ".join(section.standard_refs))

                # Content
                doc.add_paragraph(section.content)

                # Evidence references
                if section.evidence_refs:
                    ev_para = doc.add_paragraph()
                    ev_para.add_run("Supporting Evidence: ").italic = True
                    ev_para.add_run(", ".join(section.evidence_refs))

                doc.add_paragraph()  # Spacing

            elif section.section_type == PacketSectionType.EVIDENCE_INDEX:
                doc.add_heading("Evidence Index", level=1)

                # Parse markdown table and create Word table
                lines = section.content.split("\n")
                if len(lines) > 2:  # Has content beyond header
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Standard"
                    hdr_cells[1].text = "Evidence Documents"

                    for line in lines[2:]:  # Skip header and separator
                        if "|" in line:
                            parts = [p.strip() for p in line.split("|")]
                            if len(parts) >= 2:
                                row_cells = table.add_row().cells
                                row_cells[0].text = parts[0]
                                row_cells[1].text = parts[1]

                doc.add_page_break()

            elif section.section_type == PacketSectionType.EXHIBIT_LIST:
                doc.add_heading("Exhibit List", level=1)
                doc.add_paragraph(section.content)

        # Add exhibit list if we have exhibits
        if packet.exhibits:
            doc.add_heading("Exhibits", level=1)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Exhibit #"
            hdr_cells[1].text = "Title"
            hdr_cells[2].text = "Description"
            hdr_cells[3].text = "Standards"

            for exhibit in packet.exhibits:
                if exhibit.included:
                    row_cells = table.add_row().cells
                    row_cells[0].text = exhibit.exhibit_number
                    row_cells[1].text = exhibit.title
                    row_cells[2].text = exhibit.description[:100]
                    row_cells[3].text = ", ".join(exhibit.standard_refs[:3])

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Save to workspace
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename = f"submissions/{packet.id}_{timestamp}.docx"

        if self.workspace_manager:
            self.workspace_manager.save_file(
                packet.institution_id,
                filename,
                doc_bytes.read(),
                create_version=False
            )
            packet.docx_path = filename

        packet.status = PacketStatus.EXPORTED
        packet.update_stats()

        return {
            "success": True,
            "path": filename,
            "sections_included": len(sections),
            "exhibits_included": len([e for e in packet.exhibits if e.included]),
            "message": f"Exported packet to {filename}",
        }

    def _tool_export_zip(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Export complete submission as ZIP folder."""
        packet_id = params.get("packet_id", "")

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        if not self.workspace_manager:
            return {"error": "Workspace manager not available"}

        # First export DOCX
        docx_result = self._tool_export_docx({"packet_id": packet_id})
        if "error" in docx_result:
            return docx_result

        # Create ZIP in memory
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            # Add main document
            if packet.docx_path:
                docx_data = self.workspace_manager.load_file(
                    packet.institution_id, packet.docx_path
                )
                if docx_data:
                    zf.writestr("submission_packet.docx", docx_data)

            # Add exhibits in organized folders
            for exhibit in packet.exhibits:
                if exhibit.included and exhibit.file_path:
                    # Try to load exhibit file
                    exhibit_data = self.workspace_manager.load_file(
                        packet.institution_id, exhibit.file_path
                    )
                    if exhibit_data:
                        ext = exhibit.file_type or "pdf"
                        exhibit_filename = f"exhibits/{exhibit.exhibit_number}_{exhibit.title[:30]}.{ext}"
                        # Clean filename
                        exhibit_filename = "".join(c for c in exhibit_filename if c.isalnum() or c in "._-/")
                        zf.writestr(exhibit_filename, exhibit_data)

            # Add manifest
            manifest = {
                "packet_id": packet.id,
                "name": packet.name,
                "submission_type": packet.submission_type.value,
                "accrediting_body": packet.accrediting_body,
                "created_at": packet.created_at,
                "exported_at": now_iso(),
                "sections": len(packet.sections),
                "exhibits": len([e for e in packet.exhibits if e.included]),
                "standards_covered": packet.standards_covered,
            }
            zf.writestr("manifest.json", json.dumps(manifest, indent=2))

        zip_buffer.seek(0)

        # Save ZIP
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        zip_filename = f"submissions/{packet.id}_{timestamp}.zip"

        self.workspace_manager.save_file(
            packet.institution_id,
            zip_filename,
            zip_buffer.read(),
            create_version=False
        )
        packet.zip_path = zip_filename

        return {
            "success": True,
            "path": zip_filename,
            "includes_docx": bool(packet.docx_path),
            "exhibits_included": len([e for e in packet.exhibits if e.included]),
            "message": f"Exported complete submission to {zip_filename}",
        }

    def _tool_get_summary(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Get packet summary statistics."""
        packet_id = params.get("packet_id", "")

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        packet.update_stats()

        return {
            "packet_id": packet.id,
            "name": packet.name,
            "status": packet.status.value,
            "submission_type": packet.submission_type.value,
            "accrediting_body": packet.accrediting_body,
            "total_sections": packet.total_sections,
            "sections_approved": packet.sections_approved,
            "total_exhibits": packet.total_exhibits,
            "standards_covered": len(packet.standards_covered),
            "findings_addressed": packet.findings_addressed,
            "is_valid": packet.is_valid,
            "validation_issues": len(packet.validation_issues),
            "docx_path": packet.docx_path,
            "zip_path": packet.zip_path,
        }

    def _tool_save(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Save packet to workspace."""
        packet_id = params.get("packet_id", "")

        packet = self._get_packet(packet_id)
        if not packet:
            return {"error": f"Packet not found: {packet_id}"}

        if not self.workspace_manager:
            return {"error": "Workspace manager not available"}

        packet.update_stats()

        filename = f"submissions/{packet.id}.json"
        self.workspace_manager.save_file(
            packet.institution_id,
            filename,
            packet.to_dict()
        )

        return {
            "success": True,
            "packet_id": packet.id,
            "path": filename,
            "message": f"Saved packet to {filename}",
        }

    # Workflow methods

    def assemble_packet(
        self,
        institution_id: str,
        name: str,
        accrediting_body: str,
        findings_report_id: str,
        submission_type: str = "response_to_findings",
    ) -> Dict[str, Any]:
        """Assemble a complete packet from findings."""
        # Create packet
        result = self._tool_create_packet({
            "institution_id": institution_id,
            "name": name,
            "accrediting_body": accrediting_body,
            "submission_type": submission_type,
        })
        if "error" in result:
            return result

        packet_id = result["packet_id"]

        # Load findings
        result = self._tool_load_findings({
            "packet_id": packet_id,
            "findings_report_id": findings_report_id,
        })
        if "error" in result:
            return result

        # Build evidence index
        result = self._tool_build_evidence_index({"packet_id": packet_id})
        if "error" in result:
            return result

        # Save
        result = self._tool_save({"packet_id": packet_id})
        if "error" in result:
            return result

        return self._tool_get_summary({"packet_id": packet_id})

    def run_workflow(self, action: str, inputs: Dict[str, Any]) -> AgentResult:
        """Run packet workflow."""
        if action == "assemble_packet":
            result = self.assemble_packet(
                institution_id=inputs.get("institution_id", ""),
                name=inputs.get("name", ""),
                accrediting_body=inputs.get("accrediting_body", ""),
                findings_report_id=inputs.get("findings_report_id", ""),
                submission_type=inputs.get("submission_type", "response_to_findings"),
            )
            if "error" in result:
                return AgentResult.error(result["error"])
            return AgentResult.success(data=result, confidence=0.8)

        return AgentResult.error(f"Unknown workflow: {action}")
