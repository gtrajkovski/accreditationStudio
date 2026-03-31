"""Consulting Mode API.

Provides consulting-level deliverables that replace $150-300/hr consultants:
- Readiness assessments
- Pre-visit checklists
- Guided self-assessments

Endpoints:
- GET /api/consulting/readiness-assessment/<id> - Generate assessment
- POST /api/consulting/readiness-assessment/<id>/export - Export as PDF
- GET /api/consulting/pre-visit-checklist/<id> - Generate checklist
- POST /api/consulting/pre-visit-checklist/<id>/export - Export as DOCX
- GET /api/consulting/self-assessment/<id> - Get sections list
- GET /api/consulting/self-assessment/<id>/<section> - Get section detail
- POST /api/consulting/self-assessment/<id>/<section>/complete - Mark reviewed
"""

from flask import Blueprint, jsonify, request, send_file
from typing import Optional
import io
import json

from src.services.consulting_service import (
    generate_readiness_assessment,
    generate_pre_visit_checklist,
    get_self_assessment_questions,
    get_self_assessment_with_ai,
    ACCSC_SECTIONS,
)

consulting_bp = Blueprint('consulting', __name__, url_prefix='/api/consulting')

_workspace_manager = None
_ai_client = None


def init_consulting_bp(workspace_manager, ai_client=None):
    """Initialize the consulting blueprint with dependencies."""
    global _workspace_manager, _ai_client
    _workspace_manager = workspace_manager
    _ai_client = ai_client


def _get_accreditor_code(institution_id: str) -> str:
    """Get accreditor code for an institution."""
    if _workspace_manager:
        inst = _workspace_manager.load_institution(institution_id)
        if inst and hasattr(inst, 'accrediting_body'):
            return inst.accrediting_body.value
    return "ACCSC"  # Default


# =============================================================================
# Readiness Assessment Endpoints
# =============================================================================

@consulting_bp.route('/readiness-assessment/<institution_id>', methods=['GET'])
def get_readiness_assessment(institution_id: str):
    """Generate comprehensive readiness assessment.

    Query params:
        accreditor: Accreditor code (default: institution's primary)

    Returns:
        {
            "success": true,
            "assessment": {
                "institution_id": "inst_xxx",
                "institution_name": "ABC Career Institute",
                "accreditor_code": "ACCSC",
                "overall_rating": "conditionally_ready",
                "readiness_score": 78,
                "sections": [
                    {
                        "section": "Administration & Management",
                        "section_code": "admin",
                        "rating": "compliant",
                        "score": 92,
                        "total_standards": 15,
                        "compliant_count": 14,
                        "partial_count": 1,
                        "non_compliant_count": 0,
                        "critical_gaps": [],
                        "findings_summary": "14 compliant, 1 partially compliant"
                    },
                    ...
                ],
                "critical_gaps": [...],
                "timeline_recommendation": "...",
                "remediation_effort": "medium",
                "executive_summary": "...",
                "computed_at": "2026-03-31T..."
            }
        }
    """
    try:
        accreditor = request.args.get('accreditor') or _get_accreditor_code(institution_id)

        assessment = generate_readiness_assessment(institution_id, accreditor)

        return jsonify({
            "success": True,
            "assessment": assessment.to_dict(),
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


@consulting_bp.route('/readiness-assessment/<institution_id>/export', methods=['POST'])
def export_readiness_assessment(institution_id: str):
    """Export readiness assessment as PDF.

    Body:
        {
            "format": "pdf",
            "accreditor": "ACCSC"
        }

    Returns:
        PDF file download
    """
    try:
        data = request.get_json() or {}
        accreditor = data.get('accreditor') or _get_accreditor_code(institution_id)

        # Generate assessment
        assessment = generate_readiness_assessment(institution_id, accreditor)

        # Generate PDF
        pdf_bytes = _generate_assessment_pdf(assessment)

        # Send as download
        filename = f"readiness_assessment_{institution_id}.pdf"
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


# =============================================================================
# Pre-Visit Checklist Endpoints
# =============================================================================

@consulting_bp.route('/pre-visit-checklist/<institution_id>', methods=['GET'])
def get_pre_visit_checklist(institution_id: str):
    """Generate pre-visit checklist.

    Query params:
        accreditor: Accreditor code (default: institution's primary)

    Returns:
        {
            "success": true,
            "checklist": {
                "institution_id": "inst_xxx",
                "accreditor_code": "ACCSC",
                "sections": {
                    "admin": [
                        {
                            "requirement": "...",
                            "section": "Administration & Management",
                            "section_code": "admin",
                            "status": "met",
                            "evidence_reference": "Pages: 12, 15",
                            "action_needed": null,
                            "standard_code": "I.A.1",
                            "page_reference": "Pages: 12, 15"
                        },
                        ...
                    ],
                    ...
                },
                "section_progress": {
                    "admin": {"met": 12, "partial": 2, "not_met": 1, "total": 15},
                    ...
                },
                "overall_progress": {
                    "met": 54,
                    "partial": 8,
                    "not_met": 3,
                    "total": 65,
                    "percent_complete": 83
                },
                "generated_at": "2026-03-31T..."
            }
        }
    """
    try:
        accreditor = request.args.get('accreditor') or _get_accreditor_code(institution_id)

        checklist = generate_pre_visit_checklist(institution_id, accreditor)

        return jsonify({
            "success": True,
            "checklist": checklist.to_dict(),
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


@consulting_bp.route('/pre-visit-checklist/<institution_id>/export', methods=['POST'])
def export_pre_visit_checklist(institution_id: str):
    """Export pre-visit checklist as DOCX.

    Body:
        {
            "format": "docx",
            "accreditor": "ACCSC"
        }

    Returns:
        DOCX file download
    """
    try:
        data = request.get_json() or {}
        accreditor = data.get('accreditor') or _get_accreditor_code(institution_id)

        # Generate checklist
        checklist = generate_pre_visit_checklist(institution_id, accreditor)

        # Generate DOCX
        docx_bytes = _generate_checklist_docx(checklist)

        # Send as download
        filename = f"pre_visit_checklist_{institution_id}.docx"
        return send_file(
            io.BytesIO(docx_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


# =============================================================================
# Self-Assessment Endpoints
# =============================================================================

@consulting_bp.route('/self-assessment/<institution_id>', methods=['GET'])
def get_self_assessment_sections(institution_id: str):
    """Get list of self-assessment sections.

    Query params:
        accreditor: Accreditor code (default: institution's primary)

    Returns:
        {
            "success": true,
            "sections": [
                {
                    "code": "admin",
                    "name": "Administration & Management",
                    "order": 1,
                    "question_count": 15,
                    "completed": false
                },
                ...
            ]
        }
    """
    try:
        accreditor = request.args.get('accreditor') or _get_accreditor_code(institution_id)

        # Get all questions to count per section
        questions = get_self_assessment_questions(accreditor)

        # Group by section
        section_counts = {}
        for q in questions:
            # Find section code from question section name
            for code, info in ACCSC_SECTIONS.items():
                if info["name"] == q.section:
                    section_counts[code] = section_counts.get(code, 0) + 1
                    break

        sections = []
        for code, info in sorted(ACCSC_SECTIONS.items(), key=lambda x: x[1]["order"]):
            sections.append({
                "code": code,
                "name": info["name"],
                "order": info["order"],
                "question_count": section_counts.get(code, 0),
                "completed": False,  # TODO: Track completion in database
            })

        return jsonify({
            "success": True,
            "sections": sections,
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


@consulting_bp.route('/self-assessment/<institution_id>/<section>', methods=['GET'])
def get_self_assessment_section(institution_id: str, section: str):
    """Get self-assessment questions for a specific section.

    Query params:
        accreditor: Accreditor code (default: institution's primary)
        with_ai: Include AI assessments (default: true)

    Returns:
        {
            "success": true,
            "section": {
                "code": "admin",
                "name": "Administration & Management"
            },
            "questions": [
                {
                    "standard_code": "I.A.1",
                    "section": "Administration & Management",
                    "requirement_text": "...",
                    "what_to_look_for": "...",
                    "evidence_to_prepare": [...],
                    "common_deficiencies": [...],
                    "ai_assessment": "Current policy meets requirements...",
                    "ai_assessment_status": "compliant"
                },
                ...
            ]
        }
    """
    try:
        accreditor = request.args.get('accreditor') or _get_accreditor_code(institution_id)
        with_ai = request.args.get('with_ai', 'true').lower() == 'true'

        # Get questions for this section
        if with_ai:
            questions = get_self_assessment_with_ai(institution_id, accreditor, section)
        else:
            questions = get_self_assessment_questions(accreditor, section)

        # Get section info
        section_info = ACCSC_SECTIONS.get(section, {})

        return jsonify({
            "success": True,
            "section": {
                "code": section,
                "name": section_info.get("name", section),
            },
            "questions": [q.to_dict() for q in questions],
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


@consulting_bp.route('/self-assessment/<institution_id>/<section>/complete', methods=['POST'])
def mark_section_complete(institution_id: str, section: str):
    """Mark self-assessment section as reviewed.

    Body:
        {
            "reviewed_by": "user_xxx",
            "notes": "Section reviewed, all items addressed"
        }

    Returns:
        {
            "success": true,
            "message": "Section marked as complete"
        }
    """
    try:
        data = request.get_json() or {}
        reviewed_by = data.get('reviewed_by')
        notes = data.get('notes')

        # TODO: Store completion in database (Phase 47)
        # For now, return success

        return jsonify({
            "success": True,
            "message": f"Section {section} marked as complete",
            "reviewed_by": reviewed_by,
            "notes": notes,
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
        }), 500


# =============================================================================
# PDF/DOCX Export Helpers
# =============================================================================

def _generate_assessment_pdf(assessment) -> bytes:
    """Generate PDF for readiness assessment.

    Uses WeasyPrint for PDF generation.
    """
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
    except ImportError:
        # WeasyPrint not installed - return placeholder
        return b"PDF generation requires WeasyPrint. Install with: pip install weasyprint"

    # Generate HTML
    html = _generate_assessment_html(assessment)

    # Generate PDF
    font_config = FontConfiguration()
    pdf = HTML(string=html).write_pdf(font_config=font_config)

    return pdf


def _generate_assessment_html(assessment) -> str:
    """Generate HTML for assessment PDF."""
    rating_colors = {
        "ready": "#4ade80",
        "conditionally_ready": "#fbbf24",
        "not_ready": "#ef4444",
    }

    rating_labels = {
        "ready": "READY",
        "conditionally_ready": "CONDITIONALLY READY",
        "not_ready": "NOT READY",
    }

    color = rating_colors.get(assessment.overall_rating, "#9ca3af")
    label = rating_labels.get(assessment.overall_rating, assessment.overall_rating.upper())

    sections_html = ""
    for section in assessment.sections:
        gaps_html = ""
        if section.critical_gaps:
            gaps_html = "<ul>"
            for gap in section.critical_gaps[:3]:
                gaps_html += f"<li><strong>{gap['severity'].upper()}:</strong> {gap['summary']}</li>"
            gaps_html += "</ul>"

        sections_html += f"""
        <div class="section">
            <h3>{section.section} - Score: {section.score}/100</h3>
            <p><strong>Rating:</strong> {section.rating}</p>
            <p><strong>Findings:</strong> {section.findings_summary}</p>
            {gaps_html if gaps_html else '<p>No critical gaps.</p>'}
        </div>
        """

    critical_gaps_html = ""
    if assessment.critical_gaps:
        critical_gaps_html = "<h2>Critical Gaps Requiring Immediate Attention</h2><ul>"
        for gap in assessment.critical_gaps[:10]:
            critical_gaps_html += f"<li><strong>{gap['severity'].upper()}:</strong> {gap['message']}</li>"
        critical_gaps_html += "</ul>"

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Readiness Assessment</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #1a1a2e; }}
            h2 {{ color: #16213e; margin-top: 30px; }}
            h3 {{ color: #0f3460; }}
            .cover {{ text-align: center; page-break-after: always; }}
            .rating {{ font-size: 48px; font-weight: bold; color: {color}; margin: 20px 0; }}
            .score {{ font-size: 36px; margin: 10px 0; }}
            .section {{ margin: 20px 0; padding: 15px; border-left: 4px solid #0f3460; background: #f9f9f9; }}
            .footer {{ margin-top: 50px; font-size: 12px; color: #666; }}
            ul {{ line-height: 1.8; }}
        </style>
    </head>
    <body>
        <div class="cover">
            <h1>Accreditation Readiness Assessment</h1>
            <p><strong>{assessment.institution_name}</strong></p>
            <p>Accreditor: {assessment.accreditor_code}</p>
            <p>Date: {assessment.computed_at[:10]}</p>
            <div class="rating">{label}</div>
            <div class="score">Readiness Score: {assessment.readiness_score}/100</div>
        </div>

        <h2>Executive Summary</h2>
        <p style="line-height: 1.6;">{assessment.executive_summary}</p>

        <h2>Timeline Recommendation</h2>
        <p style="line-height: 1.6;">{assessment.timeline_recommendation}</p>

        <p><strong>Estimated Remediation Effort:</strong> {assessment.remediation_effort.upper()}</p>

        <h2 style="page-break-before: always;">Section-by-Section Analysis</h2>
        {sections_html}

        {critical_gaps_html}

        <div class="footer">
            <p>Generated by AccreditAI Consulting Mode</p>
            <p>This assessment is based on current documentation and audit findings.</p>
        </div>
    </body>
    </html>
    """

    return html


def _generate_checklist_docx(checklist) -> bytes:
    """Generate DOCX for pre-visit checklist.

    Uses python-docx for DOCX generation.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # python-docx not installed - return placeholder
        return b"DOCX generation requires python-docx. Install with: pip install python-docx"

    doc = Document()

    # Title
    title = doc.add_heading('Pre-Visit Checklist', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f'Institution: {checklist.institution_id}')
    doc.add_paragraph(f'Accreditor: {checklist.accreditor_code}')
    doc.add_paragraph(f'Generated: {checklist.generated_at[:10]}')

    # Overall progress
    doc.add_heading('Overall Progress', 1)
    progress = checklist.overall_progress
    doc.add_paragraph(
        f"Complete: {progress['met']}/{progress['total']} ({progress['percent_complete']}%)"
    )
    doc.add_paragraph(f"Partial: {progress['partial']}")
    doc.add_paragraph(f"Not Met: {progress['not_met']}")

    # Sections
    for section_code, items in sorted(
        checklist.sections.items(),
        key=lambda x: ACCSC_SECTIONS[x[0]]["order"]
    ):
        section_info = ACCSC_SECTIONS[section_code]
        doc.add_page_break()
        doc.add_heading(section_info["name"], 1)

        # Section progress
        prog = checklist.section_progress[section_code]
        doc.add_paragraph(
            f"Progress: {prog['met']}/{prog['total']} items complete"
        )

        # Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Requirement'
        hdr_cells[1].text = 'Status'
        hdr_cells[2].text = 'Evidence'
        hdr_cells[3].text = 'Action Needed'

        # Data rows
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.requirement[:100]  # Truncate long text
            row_cells[1].text = item.status.upper()
            row_cells[2].text = item.evidence_reference or "N/A"
            row_cells[3].text = item.action_needed or "None"

    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph('Generated by AccreditAI Consulting Mode')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
