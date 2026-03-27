"""Checklist API endpoints.

Provides REST API for managing accreditation checklists:
- Create and auto-fill checklists from audit findings
- View and edit checklist responses
- Export checklists to DOCX
"""

import json
import logging
from flask import Blueprint, request, jsonify, Response
from typing import Optional

logger = logging.getLogger(__name__)

from src.core.models import (
    AgentSession,
    SessionStatus,
    FilledChecklist,
    ChecklistResponse,
    ChecklistResponseStatus,
    ComplianceStatus,
    now_iso,
    generate_id,
)


checklists_bp = Blueprint("checklists", __name__, url_prefix="/api/institutions")
_workspace_manager = None


def init_checklists_bp(workspace_manager):
    """Initialize blueprint with dependencies."""
    global _workspace_manager
    _workspace_manager = workspace_manager


@checklists_bp.route("/<institution_id>/checklists", methods=["GET"])
def list_checklists(institution_id: str):
    """List all checklists for an institution."""
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    institution = _workspace_manager.load_institution(institution_id)
    if not institution:
        return jsonify({"error": "Institution not found"}), 404

    # List checklists from workspace
    checklists = []

    # Check institution-level checklists
    try:
        checklist_files = _workspace_manager.list_files(
            institution_id, "checklists", pattern="*.json"
        )
        for file_info in checklist_files:
            data = _workspace_manager.load_file(institution_id, f"checklists/{file_info['name']}")
            if data:
                checklists.append({
                    "id": data.get("id"),
                    "name": data.get("name"),
                    "accrediting_body": data.get("accrediting_body"),
                    "status": data.get("status"),
                    "total_items": data.get("total_items", 0),
                    "items_completed": data.get("items_completed", 0),
                    "items_compliant": data.get("items_compliant", 0),
                    "items_non_compliant": data.get("items_non_compliant", 0),
                    "created_at": data.get("created_at"),
                    "updated_at": data.get("updated_at"),
                })
    except Exception as e:
        logger.debug("No checklists found for institution %s: %s", institution_id, e)

    # Check program-level checklists
    for program in institution.programs:
        try:
            checklist_files = _workspace_manager.list_files(
                institution_id, f"programs/{program.id}/checklists", pattern="*.json"
            )
            for file_info in checklist_files:
                data = _workspace_manager.load_file(
                    institution_id, f"programs/{program.id}/checklists/{file_info['name']}"
                )
                if data:
                    checklists.append({
                        "id": data.get("id"),
                        "name": data.get("name"),
                        "program_id": program.id,
                        "program_name": program.name,
                        "accrediting_body": data.get("accrediting_body"),
                        "status": data.get("status"),
                        "total_items": data.get("total_items", 0),
                        "items_completed": data.get("items_completed", 0),
                        "items_compliant": data.get("items_compliant", 0),
                        "items_non_compliant": data.get("items_non_compliant", 0),
                        "created_at": data.get("created_at"),
                        "updated_at": data.get("updated_at"),
                    })
        except Exception as e:
            logger.debug("Failed to load checklists for program %s: %s", program.id, e)

    # Sort by updated_at descending
    checklists.sort(key=lambda x: x.get("updated_at", ""), reverse=True)

    return jsonify({"checklists": checklists})


@checklists_bp.route("/<institution_id>/checklists", methods=["POST"])
def create_checklist(institution_id: str):
    """Create and auto-fill a new checklist.

    Body:
    {
        "standards_library_id": "std_accsc",
        "audit_ids": ["audit_xxx"],  // optional
        "program_id": "",  // optional
        "name": "My Checklist"  // optional
    }
    """
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    institution = _workspace_manager.load_institution(institution_id)
    if not institution:
        return jsonify({"error": "Institution not found"}), 404

    data = request.get_json() or {}
    standards_library_id = data.get("standards_library_id")

    if not standards_library_id:
        return jsonify({"error": "standards_library_id required"}), 400

    # Create agent session
    session = AgentSession(
        agent_type="checklist",
        institution_id=institution_id,
        status=SessionStatus.RUNNING,
    )

    # Run auto-fill
    try:
        from src.agents.checklist_agent import ChecklistAgent

        agent = ChecklistAgent(session, _workspace_manager)
        result = agent.run_auto_fill(
            institution_id=institution_id,
            standards_library_id=standards_library_id,
            audit_ids=data.get("audit_ids"),
            program_id=data.get("program_id", ""),
            name=data.get("name", ""),
        )

        if "error" in result:
            return jsonify(result), 400

        return jsonify({
            "success": True,
            "checklist_id": result.get("checklist_id"),
            "name": result.get("name"),
            "total_items": result.get("total_items"),
            "items_completed": result.get("items_completed"),
            "items_compliant": result.get("items_compliant"),
            "items_non_compliant": result.get("items_non_compliant"),
            "items_needs_review": result.get("items_needs_review"),
            "completion_rate": result.get("completion_rate"),
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@checklists_bp.route("/<institution_id>/checklists/<checklist_id>", methods=["GET"])
def get_checklist(institution_id: str, checklist_id: str):
    """Get a filled checklist with all responses."""
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    # Try to find the checklist
    checklist_data = None

    # Check institution-level
    checklist_data = _workspace_manager.load_file(
        institution_id, f"checklists/{checklist_id}.json"
    )

    # Check program-level if not found
    if not checklist_data:
        institution = _workspace_manager.load_institution(institution_id)
        if institution:
            for program in institution.programs:
                checklist_data = _workspace_manager.load_file(
                    institution_id, f"programs/{program.id}/checklists/{checklist_id}.json"
                )
                if checklist_data:
                    break

    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    return jsonify(checklist_data)


@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/items/<item_number>",
    methods=["PUT"]
)
def update_checklist_item(institution_id: str, checklist_id: str, item_number: str):
    """Update a specific checklist item response.

    Body:
    {
        "compliance_status": "compliant",
        "narrative_response": "...",
        "evidence_summary": "...",
        "human_notes": "...",
        "approve": true
    }
    """
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    data = request.get_json() or {}

    # Find and load checklist
    checklist_data = None
    checklist_path = None

    # Check institution-level
    checklist_data = _workspace_manager.load_file(
        institution_id, f"checklists/{checklist_id}.json"
    )
    if checklist_data:
        checklist_path = f"checklists/{checklist_id}.json"

    # Check program-level if not found
    if not checklist_data:
        institution = _workspace_manager.load_institution(institution_id)
        if institution:
            for program in institution.programs:
                checklist_data = _workspace_manager.load_file(
                    institution_id, f"programs/{program.id}/checklists/{checklist_id}.json"
                )
                if checklist_data:
                    checklist_path = f"programs/{program.id}/checklists/{checklist_id}.json"
                    break

    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    # Find the item
    checklist = FilledChecklist.from_dict(checklist_data)
    response = next(
        (r for r in checklist.responses if r.item_number == item_number),
        None
    )

    if not response:
        return jsonify({"error": f"Item not found: {item_number}"}), 404

    # Update fields
    if "compliance_status" in data:
        response.compliance_status = ComplianceStatus(data["compliance_status"])
    if "narrative_response" in data:
        response.narrative_response = data["narrative_response"]
    if "evidence_summary" in data:
        response.evidence_summary = data["evidence_summary"]
    if "human_notes" in data:
        response.human_notes = data["human_notes"]

    if data.get("approve"):
        response.response_status = ChecklistResponseStatus.APPROVED
    else:
        response.response_status = ChecklistResponseStatus.HUMAN_EDITED

    response.last_updated = now_iso()

    # Update stats and save
    checklist.update_stats()
    checklist.updated_at = now_iso()

    _workspace_manager.save_file(institution_id, checklist_path, checklist.to_dict())

    return jsonify({
        "success": True,
        "item_number": item_number,
        "response_status": response.response_status.value,
        "compliance_status": response.compliance_status.value,
    })


@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/export",
    methods=["GET"]
)
def export_checklist(institution_id: str, checklist_id: str):
    """Export checklist to DOCX format."""
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    # Find and load checklist
    checklist_data = _workspace_manager.load_file(
        institution_id, f"checklists/{checklist_id}.json"
    )

    if not checklist_data:
        institution = _workspace_manager.load_institution(institution_id)
        if institution:
            for program in institution.programs:
                checklist_data = _workspace_manager.load_file(
                    institution_id, f"programs/{program.id}/checklists/{checklist_id}.json"
                )
                if checklist_data:
                    break

    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    checklist = FilledChecklist.from_dict(checklist_data)

    # Generate DOCX
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        # Title
        title = doc.add_heading(checklist.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        doc.add_heading("Summary", level=1)
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = "Table Grid"

        rows = [
            ("Total Items", str(checklist.total_items)),
            ("Completed", str(checklist.items_completed)),
            ("Compliant", str(checklist.items_compliant)),
            ("Partial", str(checklist.items_partial)),
            ("Non-Compliant", str(checklist.items_non_compliant)),
        ]

        for i, (label, value) in enumerate(rows):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Group responses by category
        by_category = {}
        for response in checklist.responses:
            cat = response.category or "Other"
            if cat not in by_category:
                by_category[cat] = []
            by_category[cat].append(response)

        # Add each category
        for category, responses in sorted(by_category.items()):
            doc.add_heading(category, level=1)

            for response in responses:
                # Item header
                para = doc.add_paragraph()
                run = para.add_run(f"{response.item_number}: {response.item_description}")
                run.bold = True

                # Status
                status_text = response.compliance_status.value.replace("_", " ").title()
                status_para = doc.add_paragraph()
                status_para.add_run("Status: ").bold = True
                status_run = status_para.add_run(status_text)

                if response.compliance_status == ComplianceStatus.COMPLIANT:
                    status_run.font.color.rgb = RGBColor(0, 128, 0)
                elif response.compliance_status == ComplianceStatus.NON_COMPLIANT:
                    status_run.font.color.rgb = RGBColor(255, 0, 0)
                elif response.compliance_status == ComplianceStatus.PARTIAL:
                    status_run.font.color.rgb = RGBColor(255, 165, 0)

                # Narrative
                if response.narrative_response:
                    doc.add_paragraph(response.narrative_response)

                # Evidence
                if response.evidence_summary:
                    evidence_para = doc.add_paragraph()
                    evidence_para.add_run("Evidence: ").bold = True
                    evidence_para.add_run(response.evidence_summary[:500])

                doc.add_paragraph()  # Spacing

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{checklist.name.replace(' ', '_')}.docx"

        return Response(
            buffer.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/approve-all",
    methods=["POST"]
)
def approve_all_items(institution_id: str, checklist_id: str):
    """Approve all auto-filled items in a checklist."""
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    # Find and load checklist
    checklist_data = None
    checklist_path = None

    checklist_data = _workspace_manager.load_file(
        institution_id, f"checklists/{checklist_id}.json"
    )
    if checklist_data:
        checklist_path = f"checklists/{checklist_id}.json"

    if not checklist_data:
        institution = _workspace_manager.load_institution(institution_id)
        if institution:
            for program in institution.programs:
                checklist_data = _workspace_manager.load_file(
                    institution_id, f"programs/{program.id}/checklists/{checklist_id}.json"
                )
                if checklist_data:
                    checklist_path = f"programs/{program.id}/checklists/{checklist_id}.json"
                    break

    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    checklist = FilledChecklist.from_dict(checklist_data)
    approved_count = 0

    for response in checklist.responses:
        if response.response_status in (
            ChecklistResponseStatus.AUTO_FILLED,
            ChecklistResponseStatus.HUMAN_EDITED,
        ):
            response.response_status = ChecklistResponseStatus.APPROVED
            response.last_updated = now_iso()
            approved_count += 1

    checklist.update_stats()
    checklist.updated_at = now_iso()

    _workspace_manager.save_file(institution_id, checklist_path, checklist.to_dict())

    return jsonify({
        "success": True,
        "approved_count": approved_count,
        "total_items": checklist.total_items,
    })


# =============================================================================
# Phase 7 Enhanced Endpoints
# =============================================================================

@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/validate",
    methods=["POST"]
)
def validate_checklist(institution_id: str, checklist_id: str):
    """Validate checklist claims against actual document content.

    Body:
    {
        "item_numbers": ["1.1", "1.2"],  // optional, empty for all
        "strict_mode": false  // optional
    }
    """
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    data = request.get_json() or {}

    # Load checklist and create agent
    checklist_data, checklist_path = _find_checklist(institution_id, checklist_id)
    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    try:
        from src.agents.checklist_agent import ChecklistAgent

        session = AgentSession(
            agent_type="checklist",
            institution_id=institution_id,
            status=SessionStatus.RUNNING,
        )

        agent = ChecklistAgent(session, _workspace_manager)

        # Load checklist into agent
        agent._current_checklist = FilledChecklist.from_dict(checklist_data)

        # Run validation
        result = agent._tool_validate_against_documents({
            "item_numbers": data.get("item_numbers", []),
            "strict_mode": data.get("strict_mode", False),
        })

        if "error" in result:
            return jsonify(result), 400

        # Save updated checklist
        if agent._current_checklist:
            _workspace_manager.save_file(
                institution_id,
                checklist_path,
                agent._current_checklist.to_dict()
            )

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/page-references",
    methods=["POST"]
)
def generate_page_references(institution_id: str, checklist_id: str):
    """Generate page references for checklist evidence.

    Body:
    {
        "item_numbers": ["1.1", "1.2"],  // optional, empty for all
        "include_excerpts": true  // optional
    }
    """
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    data = request.get_json() or {}

    checklist_data, checklist_path = _find_checklist(institution_id, checklist_id)
    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    try:
        from src.agents.checklist_agent import ChecklistAgent

        session = AgentSession(
            agent_type="checklist",
            institution_id=institution_id,
            status=SessionStatus.RUNNING,
        )

        agent = ChecklistAgent(session, _workspace_manager)
        agent._current_checklist = FilledChecklist.from_dict(checklist_data)

        result = agent._tool_generate_page_references({
            "item_numbers": data.get("item_numbers", []),
            "include_excerpts": data.get("include_excerpts", True),
        })

        if "error" in result:
            return jsonify(result), 400

        # Save updated checklist
        if agent._current_checklist:
            _workspace_manager.save_file(
                institution_id,
                checklist_path,
                agent._current_checklist.to_dict()
            )

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/export-linked",
    methods=["POST"]
)
def export_with_evidence_links(institution_id: str, checklist_id: str):
    """Export checklist with hyperlinked evidence references.

    Body:
    {
        "format": "docx",  // docx, pdf, json
        "include_appendix": true,
        "group_by": "category"  // category, status, section
    }
    """
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    data = request.get_json() or {}

    checklist_data, checklist_path = _find_checklist(institution_id, checklist_id)
    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    try:
        from src.agents.checklist_agent import ChecklistAgent

        session = AgentSession(
            agent_type="checklist",
            institution_id=institution_id,
            status=SessionStatus.RUNNING,
        )

        agent = ChecklistAgent(session, _workspace_manager)
        agent._current_checklist = FilledChecklist.from_dict(checklist_data)

        result = agent._tool_export_with_evidence_links({
            "format": data.get("format", "docx"),
            "include_appendix": data.get("include_appendix", True),
            "group_by": data.get("group_by", "category"),
        })

        if "error" in result:
            return jsonify(result), 400

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@checklists_bp.route(
    "/<institution_id>/checklists/<checklist_id>/completion-status",
    methods=["GET"]
)
def check_completion_status(institution_id: str, checklist_id: str):
    """Get detailed completion progress by category.

    Query params:
    - category: specific category to check (optional)
    - include_blockers: true/false (default: true)
    """
    if not _workspace_manager:
        return jsonify({"error": "Service not initialized"}), 500

    category = request.args.get("category")
    include_blockers = request.args.get("include_blockers", "true").lower() == "true"

    checklist_data, _ = _find_checklist(institution_id, checklist_id)
    if not checklist_data:
        return jsonify({"error": "Checklist not found"}), 404

    try:
        from src.agents.checklist_agent import ChecklistAgent

        session = AgentSession(
            agent_type="checklist",
            institution_id=institution_id,
            status=SessionStatus.RUNNING,
        )

        agent = ChecklistAgent(session, _workspace_manager)
        agent._current_checklist = FilledChecklist.from_dict(checklist_data)

        result = agent._tool_check_completion_status({
            "category": category,
            "include_blockers": include_blockers,
        })

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _find_checklist(institution_id: str, checklist_id: str):
    """Helper to find a checklist in workspace.

    Returns (checklist_data, checklist_path) or (None, None) if not found.
    """
    # Check institution-level
    checklist_data = _workspace_manager.load_file(
        institution_id, f"checklists/{checklist_id}.json"
    )
    if checklist_data:
        return checklist_data, f"checklists/{checklist_id}.json"

    # Check program-level
    institution = _workspace_manager.load_institution(institution_id)
    if institution:
        for program in institution.programs:
            checklist_data = _workspace_manager.load_file(
                institution_id, f"programs/{program.id}/checklists/{checklist_id}.json"
            )
            if checklist_data:
                return checklist_data, f"programs/{program.id}/checklists/{checklist_id}.json"

    return None, None
